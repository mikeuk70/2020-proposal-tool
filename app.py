"""
20.20 Design Agency — Proposal Generator
Hosted Flask app for LawLiss / 20.20
"""

import os, json, uuid, threading, queue, time, base64, re, copy, zipfile, tempfile, shutil
import anthropic
from flask import Flask, request, jsonify, send_file, Response
from werkzeug.exceptions import RequestEntityTooLarge
from pptx_builder import build_pptx_clean
def build_docx(sections, meta):
    """Build a clean Word document from generated sections for team review and sharing."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    import tempfile

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    for lvl, sz in [('Heading 1', 18), ('Heading 2', 13)]:
        s = doc.styles[lvl]
        s.font.name = 'Arial'; s.font.size = Pt(sz); s.font.bold = True
        s.font.color.rgb = RGBColor(0x11, 0x14, 0x18)

    # Collects every [CONFIRM WITH CLIENT: ...] note found while building
    # this doc, grouped by section heading, so they can be moved to a
    # dedicated internal action page at the end instead of appearing inline
    # in client-facing text. Mirrors the same behaviour in pptx_builder.py —
    # this export path needs its own copy since it has its own text-cleaning
    # function rather than sharing pptx_builder's.
    confirm_notes = []
    _confirm_re = re.compile(r'\[CONFIRM WITH CLIENT:\s*([^\]]+)\]', re.IGNORECASE)

    def _strip_confirm_notes(text, label):
        if not text:
            return text
        for m in _confirm_re.finditer(text):
            note = m.group(1).strip()
            if note:
                confirm_notes.append((label, note))
        return _confirm_re.sub('', text)

    def _c(t):
        if not t: return ''
        t = re.sub(r'\*\*([^*]+)\*\*', r'\1', t)
        t = re.sub(r'\*([^*]+)\*', r'\1', t)
        t = re.sub(r'^#{1,4}\s*', '', t, flags=re.MULTILINE)
        return t.strip()

    def add_body(body):
        for line in _c(body).splitlines():
            s = line.strip()
            if not s: continue
            if re.match(r'^(Objective|Process|Deliverables|Meetings[^:]*|Presentations):?\s*$', s, re.I):
                doc.add_heading(s.rstrip(':'), level=2)
            elif s.startswith(('-', '\u2022', '*')) or re.match(r'^\d+[.):]', s):
                text = re.sub(r'^[-\u2022*]\s*|^\d+[.):]+\s*', '', s)
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(text); run.font.name = 'Arial'; run.font.size = Pt(11)
            else:
                p = doc.add_paragraph(); run = p.add_run(s)
                run.font.name = 'Arial'; run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(6)

    # Title block
    tp = doc.add_paragraph()
    tr = tp.add_run(meta.get('venue', 'Proposal'))
    tr.font.name = 'Arial'; tr.font.size = Pt(26); tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x11, 0x14, 0x18)

    sp = doc.add_paragraph()
    sr = sp.add_run('Hospitality design proposal')
    sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    details = [meta.get('client','')]
    if meta.get('contact'): details.append('Prepared for ' + meta['contact'] + (', ' + meta['role'] if meta.get('role') else ''))
    if meta.get('date'): details.append(meta['date'])
    dp = doc.add_paragraph('  |  '.join(d for d in details if d))
    if dp.runs: dp.runs[0].font.size = Pt(10); dp.runs[0].font.color.rgb = RGBColor(0x88,0x88,0x88)

    cp = doc.add_paragraph('CONFIDENTIAL  \u00a9  20.20 Limited 2026')
    if cp.runs: cp.runs[0].font.size = Pt(8.5); cp.runs[0].font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
    doc.add_page_break()

    for sec in sections:
        body = sec.get('body','')
        if not body.strip(): continue
        heading = sec.get('heading', sec.get('id','').replace('_',' ').title())
        body = _strip_confirm_notes(body, heading)
        doc.add_heading(heading, level=1)
        add_body(body)
        doc.add_paragraph()

    if confirm_notes:
        doc.add_page_break()
        h = doc.add_heading('Internal use only — points to confirm before sending', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0xA3, 0x2D, 0x2D)
        note_p = doc.add_paragraph()
        note_r = note_p.add_run(
            'Assumptions and gaps flagged while drafting. Resolve these with the client '
            'or remove this page before the proposal goes out.'
        )
        note_r.font.italic = True
        note_r.font.size = Pt(10)
        note_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        seen_labels = []
        by_label = {}
        for label, note in confirm_notes:
            if label not in by_label:
                by_label[label] = []
                seen_labels.append(label)
            by_label[label].append(note)
        for label in seen_labels:
            doc.add_heading(label, level=2)
            for note in by_label[label]:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(note); run.font.name = 'Arial'; run.font.size = Pt(11)

    tmp = tempfile.mkdtemp(prefix='2020_docx_')
    slug = re.sub(r'[^a-zA-Z0-9]+', '_', meta.get('venue','Proposal'))
    path = os.path.join(tmp, f'{slug}_20.20_Proposal.docx')
    doc.save(path)
    return path



app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 40 * 1024 * 1024  # 40MB max upload — main brief + up to 3 supporting PDFs. Note: this caps request SIZE in bytes, not tokens — the rate limit that actually bites on complex briefs is Anthropic's per-minute token limit, not file size.

# ── CONFIG ────────────────────────────────────────────────────────────────────
ANTHROPIC_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
# Shared secret for the /usage dashboard — set USAGE_DASHBOARD_KEY in
# Railway's environment variables. Without it, /usage refuses all access
# rather than defaulting to open, since the dashboard shows client names.
USAGE_DASHBOARD_KEY = os.environ.get('USAGE_DASHBOARD_KEY', '')
# Find template file - check several locations
_here = os.path.dirname(os.path.abspath(__file__))
_candidates = [
    os.path.join(_here, '2020_template_slim_b64.txt'),
    '/app/2020_template_slim_b64.txt',
    os.path.join(os.getcwd(), '2020_template_slim_b64.txt'),
]
TEMPLATE_PATH = next((p for p in _candidates if os.path.exists(p)), _candidates[0])

# File-based job store — survives restarts and works across gunicorn workers
JOBS_DIR = os.path.join(tempfile.gettempdir(), '2020_jobs')
os.makedirs(JOBS_DIR, exist_ok=True)

def job_path(job_id):
    return os.path.join(JOBS_DIR, f'{job_id}.json')

def pptx_path_for(job_id):
    return os.path.join(JOBS_DIR, f'{job_id}.pptx')

def load_job(job_id):
    p = job_path(job_id)
    if not os.path.exists(p):
        return None
    try:
        with open(p, 'r') as f:
            return json.load(f)
    except Exception:
        return None

def save_job(job_id, job):
    p = job_path(job_id)
    try:
        with open(p, 'w') as f:
            json.dump(job, f)
    except Exception:
        pass

def update_job(job_id, **kwargs):
    job = load_job(job_id) or {}
    job.update(kwargs)
    save_job(job_id, job)

def append_progress(job_id, msg, pct=None):
    job = load_job(job_id) or {}
    prog = job.get('progress', [])
    prog.append({'msg': msg, 'pct': pct})
    job['progress'] = prog
    save_job(job_id, job)

def append_section(job_id, section):
    job = load_job(job_id) or {}
    secs = job.get('sections', [])
    secs.append(section)
    job['sections'] = secs
    save_job(job_id, job)


# ── USAGE LOG ─────────────────────────────────────────────────────────────────
# Tracks every generation attempt so usage can be reviewed later (client,
# venue, brief type, outcome, duration). Append-only JSONL — one JSON object
# per line — so a single corrupted line never breaks the whole file, and
# writes never need to read the existing file first.
#
# IMPORTANT — persistence: JOBS_DIR lives under tempfile.gettempdir(), which
# on Railway is ephemeral storage wiped on every redeploy/restart. This log
# file lives in the same place and has the same limitation. It is NOT a
# substitute for a real database if usage data needs to survive redeploys
# long-term — treat this as a working solution to get visibility now, and
# revisit storage (e.g. a Railway volume, or writing to an external sheet/DB)
# if this data needs to be durable indefinitely.
USAGE_LOG_PATH = os.path.join(JOBS_DIR, '_usage_log.jsonl')

def log_usage_event(job_id, event, **fields):
    """Append one usage event. event is a short string like 'started',
    'completed', 'failed'. Extra fields (client, venue, brief_type, etc.)
    are merged in. Never raises — a logging failure should never break
    the actual proposal generation."""
    try:
        entry = {
            'ts': time.time(),
            'date': time.strftime('%Y-%m-%d %H:%M:%S'),
            'job_id': job_id,
            'event': event,
        }
        entry.update(fields)
        with open(USAGE_LOG_PATH, 'a') as f:
            f.write(json.dumps(entry) + '\n')
    except Exception:
        pass

def read_usage_log():
    """Read all usage events, oldest first. Skips any corrupted lines
    rather than failing the whole read."""
    events = []
    if not os.path.exists(USAGE_LOG_PATH):
        return events
    with open(USAGE_LOG_PATH, 'r') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                events.append(json.loads(line))
            except Exception:
                continue
    return events

def migrate_existing_jobs_to_usage_log():
    """Best-effort backfill: scan whatever job .json files currently exist
    in JOBS_DIR and add a 'completed'/'failed' usage event for any job not
    already represented in the usage log. Only catches jobs that survived
    up to the most recent server restart — JOBS_DIR is ephemeral, so this
    cannot recover jobs from before a redeploy that already happened.
    Returns a summary dict so the caller can report exactly what happened."""
    existing_job_ids = {e.get('job_id') for e in read_usage_log()}
    found, migrated, skipped = 0, 0, 0
    if not os.path.isdir(JOBS_DIR):
        return {'found': 0, 'migrated': 0, 'skipped': 0}
    for fname in os.listdir(JOBS_DIR):
        if not fname.endswith('.json') or fname.startswith('_'):
            continue
        job_id = fname[:-5]
        found += 1
        if job_id in existing_job_ids:
            skipped += 1
            continue
        job = load_job(job_id)
        if not job:
            skipped += 1
            continue
        meta = job.get('meta', {})
        status = job.get('status', 'unknown')
        log_usage_event(
            job_id, 'completed' if status == 'done' else status,
            client=meta.get('client', ''),
            venue=meta.get('venue', ''),
            brief_type=meta.get('brief_type', ''),
            is_riba=meta.get('is_riba', ''),
            error=job.get('error', ''),
            pptx_ready=bool(job.get('pptx_path') and os.path.exists(job.get('pptx_path', ''))),
            migrated_from_job_file=True,
        )
        migrated += 1
    return {'found': found, 'migrated': migrated, 'skipped': skipped}


# ── NAMESPACES ────────────────────────────────────────────────────────────────
P = 'http://schemas.openxmlformats.org/presentationml/2006/main'
A = 'http://schemas.openxmlformats.org/drawingml/2006/main'

import xml.etree.ElementTree as ET
for prefix, uri in [
    ('p', P), ('a', A),
    ('r', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'),
    ('a16', 'http://schemas.microsoft.com/office/drawing/2014/main'),
]:
    ET.register_namespace(prefix, uri)


# ── CLIENT COLOURS ────────────────────────────────────────────────────────────
CLUB_COLOURS = {
    'aston villa': '5C1A2E', 'villa': '5C1A2E',
    'newcastle': 'C9A84C', 'nufc': 'C9A84C',
    'brighton': '0057B8', 'bhafc': '0057B8', 'amex': '0057B8',
    'arsenal': 'EF0107',
    'liverpool': 'C8102E',
    'chelsea': '034694',
    'crystal palace': '1B458F', 'cpfc': '1B458F',
    'leeds': 'FFCD00', 'lufc': 'FFCD00',
    'sunderland': 'EB172B', 'safc': 'EB172B',
    'west ham': '7A263A',
    'manchester city': '6CABDD', 'man city': '6CABDD',
    'manchester united': 'DA291C', 'man utd': 'DA291C',
    'tottenham': '132257', 'spurs': '132257',
    'everton': '003399',
    'sheffield': 'EE2737',
    'nottingham forest': 'E53233',
    'leicester': '003090',
    'wolves': 'FDB913',
    'celtic': '00843D',
    'rangers': '0033A0',
}
DEFAULT_COLOUR = 'E97132'  # 20.20 orange

def detect_colour(client_name):
    if not client_name:
        return DEFAULT_COLOUR
    cl = client_name.lower()
    for key, colour in CLUB_COLOURS.items():
        if key in cl:
            return colour
    return DEFAULT_COLOUR


# ── TEXT HELPERS ──────────────────────────────────────────────────────────────
def clean(txt):
    if not txt:
        return ''
    txt = re.sub(r'\*\*([^*]+)\*\*', r'\1', txt)
    txt = re.sub(r'\*([^*]+)\*', r'\1', txt)
    txt = re.sub(r'^#{1,3}\s*', '', txt, flags=re.MULTILINE)
    txt = re.sub(r'\n{3,}', '\n\n', txt)
    return txt.strip()

def first_sentences(txt, n=2):
    s = clean(txt)
    return ' '.join(re.split(r'(?<=[.!?])\s+', s)[:n]).strip()

def explicit_bullets(txt, max_n=8):
    s = clean(txt)
    bullets = []
    for line in s.split('\n'):
        stripped = line.strip()
        if stripped.startswith(('-', '\u2022', '*')):
            item = re.sub(r'^[-\u2022*]\s*', '', stripped).strip()
            if len(item) > 8 and item not in bullets:
                bullets.append(item)
        elif re.match(r'^\d+[.):]', stripped):
            item = re.sub(r'^\d+[.):]+\s*', '', stripped).strip()
            if len(item) > 8 and item not in bullets:
                bullets.append(item)
    if not bullets:
        # Fall back: any line that reads like a deliverable item
        for line in s.split('\n'):
            line = line.strip()
            if len(line) > 15 and not re.search(r'stage \d|riba|objective|process', line, re.I):
                line = re.sub(r'^[-\u2022*]\s*', '', line)
                if line not in bullets:
                    bullets.append(line)
    return [b for b in bullets if b][:max_n]

def prose_only(txt, n=3):
    s = clean(txt)
    lines = s.split('\n')
    prose = []
    for line in lines:
        stripped = line.strip()
        # Stop at bullet lists
        if stripped.startswith(('-', '\u2022', '*')) or re.match(r'^\d+[.):]', stripped):
            break
        # Stop at section keyword headings
        if stripped.lower().rstrip(':') in ('deliverables', 'outputs', 'scope', 'process',
                                             'objective', 'approach', 'programme', 'fees',
                                             'next steps', 'our approach'):
            break
        # Skip lines that look like AI-generated stage headers e.g. "Stage 1: ... | RIBA Stage 2 | 1 week"
        if re.search(r'stage \d.*riba|riba.*stage \d|\|\s*\d+\s*week', stripped, re.IGNORECASE):
            continue
        # Skip lines that are just the section label repeated (e.g. "Your brief", "Cover letter")
        if len(stripped) < 40 and '.' not in stripped and ',' not in stripped:
            continue
        prose.append(stripped)
    text = ' '.join(l for l in prose if l)
    # Take first n sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return ' '.join(sentences[:n]).strip()

def find_section(sections, *keys):
    for key in keys:
        kl = key.lower()
        for sec in sections:
            h = sec.get('heading', '').lower()
            if kl in h or h in kl:
                return clean(sec.get('body', ''))
    return ''


# ── XML HELPERS ───────────────────────────────────────────────────────────────
def get_txbodies(root):
    return [e for e in root.iter() if e.tag == f'{{{P}}}txBody']

def full_text(tb):
    return ''.join(e.text for e in tb.iter() if e.tag == f'{{{A}}}t' and e.text)

def get_first_rPr(tb):
    for r in tb.iter(f'{{{A}}}r'):
        rPr = r.find(f'{{{A}}}rPr')
        if rPr is not None:
            return copy.deepcopy(rPr)
    return None

def make_rPr(tmpl=None, bold=False, colour=None):
    rPr = ET.Element(f'{{{A}}}rPr')
    rPr.set('lang', 'en-GB')
    rPr.set('dirty', '0')
    if tmpl is not None:
        for attr in ['sz', 'lang']:
            if tmpl.get(attr):
                rPr.set(attr, tmpl.get(attr))
        if not colour:
            for child in tmpl:
                if any(k in child.tag for k in ('Fill', 'latin', 'ea', 'cs')):
                    rPr.append(copy.deepcopy(child))
        else:
            for child in tmpl:
                if any(k in child.tag for k in ('latin', 'ea', 'cs')) and 'Fill' not in child.tag:
                    rPr.append(copy.deepcopy(child))
    if colour:
        sf = ET.SubElement(rPr, f'{{{A}}}solidFill')
        sc = ET.SubElement(sf, f'{{{A}}}srgbClr')
        sc.set('val', colour.upper().lstrip('#'))
    if bold:
        rPr.set('b', '1')
    return rPr

def set_text(tb, text, tmpl_rPr=None, bold=False, colour=None):
    for p in [e for e in tb if e.tag == f'{{{A}}}p']:
        tb.remove(p)
    p = ET.SubElement(tb, f'{{{A}}}p')
    r = ET.SubElement(p, f'{{{A}}}r')
    r.append(make_rPr(tmpl_rPr, bold=bold, colour=colour))
    ET.SubElement(r, f'{{{A}}}t').text = text

def set_paragraphs(tb, items, tmpl_rPr=None):
    for p in [e for e in tb if e.tag == f'{{{A}}}p']:
        tb.remove(p)
    for text, opts in items:
        p = ET.SubElement(tb, f'{{{A}}}p')
        pPr = ET.SubElement(p, f'{{{A}}}pPr')
        pPr.set('lvl', '0')
        if opts.get('bullet'):
            ET.SubElement(pPr, f'{{{A}}}buFont').set('typeface', 'Arial')
            ET.SubElement(pPr, f'{{{A}}}buChar').set('char', '\u2022')
        else:
            ET.SubElement(pPr, f'{{{A}}}buNone')
        if text:
            r = ET.SubElement(p, f'{{{A}}}r')
            r.append(make_rPr(tmpl_rPr, bold=opts.get('bold', False),
                               colour=opts.get('colour')))
            ET.SubElement(r, f'{{{A}}}t').text = text
        else:
            ET.SubElement(p, f'{{{A}}}endParaRPr').set('lang', 'en-GB')

def replace_colour(xml_str, old, new):
    old, new = old.upper(), new.upper()
    for v in [old, old.lower()]:
        xml_str = xml_str.replace(f'val="{v}"', f'val="{new}"')
    return xml_str


# ── SLIDE BUILDERS ────────────────────────────────────────────────────────────
def build_cover(root, venue, contact, role, date_s):
    r = copy.deepcopy(root)
    tbs = get_txbodies(r)
    for i, tb in enumerate(tbs):
        ft = full_text(tb)
        rPr = get_first_rPr(tb)
        if 'Concept Presentation' in ft:
            set_text(tb, venue, rPr)
        elif 'Stage 2' in ft and len(ft.strip()) < 25:
            set_text(tb, 'Hospitality design proposal', rPr)
        elif any(x in ft for x in ['June', '25th', "'25", 'th']):
            if date_s:
                set_text(tb, date_s, rPr)
    return r

def build_hello(root):
    r = copy.deepcopy(root)
    for tb in get_txbodies(r):
        if full_text(tb).strip() == 'PowerPoint Template':
            set_text(tb, '', get_first_rPr(tb))
    return r

def build_dark_divider(root, word):
    r = copy.deepcopy(root)
    for tb in get_txbodies(r):
        ft = full_text(tb).strip()
        rPr = get_first_rPr(tb)
        if ft == 'Hello':
            set_text(tb, word, rPr)
        elif ft == 'PowerPoint Template':
            set_text(tb, '', rPr)
    return r

def build_content_slide(root, section_label, title, intro, bullets):
    r = copy.deepcopy(root)
    for tb in get_txbodies(r):
        ft = full_text(tb)
        rPr = get_first_rPr(tb)
        if 'PowerPoint Template' in ft and len(ft) < 40:
            set_text(tb, section_label, rPr)
        elif 'Example header' in ft:
            set_text(tb, title, rPr)
        elif 'Lorem ipsum' in ft or ('lorem' in ft.lower() and len(ft) > 30):
            items = []
            # Clean intro — remove if it is just a short heading label
            clean_intro = intro.strip() if intro else ''
            if clean_intro and len(clean_intro) > 30:
                items.append((clean_intro, {}))
            if bullets:
                if items:
                    items.append(('', {}))
                for b in bullets:
                    items.append((b, {'bullet': True}))
            if items:
                set_paragraphs(tb, items, rPr)
    return r

def build_stage_slide(slide14_raw, section_label, stage_title, body, deliverables, accent):
    root = ET.fromstring(slide14_raw)
    tbs = get_txbodies(root)
    for i, tb in enumerate(tbs):
        rPr = get_first_rPr(tb)
        if i == 0:
            set_text(tb, section_label, rPr)
        elif i == 2:
            set_text(tb, stage_title, rPr)
        elif i == 3:
            prose = prose_only(body, 3)
            if prose:
                set_paragraphs(tb, [(prose, {})], rPr)
        elif i == 4:
            set_text(tb, 'Scope', rPr, colour=accent)
        elif i == 5:
            dl = [d for d in deliverables if d] if deliverables else []
            if not dl:
                dl = explicit_bullets(body, 8)
            items = [(d, {'bullet': True}) for d in dl[:8]]
            if items:
                set_paragraphs(tb, items, rPr)
        elif i == 7:
            set_text(tb, '[FEE: TBC — rate card required]', rPr, bold=True)
        elif i == 8:
            set_text(tb, 'Deliverables', rPr, colour=accent)
    raw = ET.tostring(root, encoding='unicode')
    return "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + raw

def build_fees_slide(root, stages, accent):
    r = copy.deepcopy(root)
    stage_keys = ['WORKSHOP & DEFINITION', 'CONCEPT DESIGN', 'DESIGN DEVELOPMENT', 'DESIGN INTENT']
    si = 0
    for tb in get_txbodies(r):
        ft = full_text(tb)
        rPr = get_first_rPr(tb)
        if 'PowerPoint Template' in ft and len(ft) < 40:
            set_text(tb, 'Our methodology', rPr)
        elif 'Summary fees and timings' in ft:
            set_text(tb, 'Summary fees and timings', rPr)
        elif any(k in ft for k in stage_keys) and si < len(stages):
            sd = stages[si]
            items = [(sd['title'], {'bold': True})]
            if sd.get('sub'):
                items.append((sd['sub'], {}))
            set_paragraphs(tb, items, rPr)
            si += 1
        elif 'TOTAL 20.20' in ft or 'AFEES' in ft:
            set_paragraphs(tb, [
                ('TOTAL 20.20 FEES', {'bold': True}),
                ('Planning and Interiors', {}),
                ('Identity and Graphics', {}),
                ('Strategy and Management', {}),
            ], rPr)
        elif ft.startswith('£') and 'Bobby' not in ft and len(ft) < 30:
            set_text(tb, '[FEE: TBC]', rPr, bold=True)
    return r


# ── PPTX BUILDER ─────────────────────────────────────────────────────────────

def reorder_presentation(unpacked):
    """
    Rewrite presentation.xml to show only our 10 proposal slides in the right order.
    The template has 32 slides — we select and reorder just the ones we need.
    """
    prs_path = os.path.join(unpacked, 'ppt', 'presentation.xml')
    with open(prs_path, encoding='utf-8') as f:
        prs_xml = f.read()

    P_NS = 'http://schemas.openxmlformats.org/presentationml/2006/main'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    root = ET.fromstring(prs_xml)
    sldIdLst = root.find(f'{{{P_NS}}}sldIdLst')
    if sldIdLst is None:
        return  # can't find slide list, leave as-is

    # Map rId -> slide element from original list
    existing = {sld.get(f'{{{R_NS}}}id'): sld for sld in sldIdLst}

    # Our desired slide order (rId, slide_num, label)
    DECK_ORDER = [
        ('rId2',  'Cover'),
        ('rId28', 'Hello — cover letter'),
        ('rId9',  'Our understanding'),
        ('rId29', 'Our methodology divider'),
        ('rId15', 'Stage 1'),
        ('rId10', 'Stage 2'),
        ('rId11', 'Stage 3'),
        ('rId12', 'Stages 4-6'),
        ('rId17', 'Fees and timings'),
        ('rId27', 'Next steps'),
    ]

    # Clear the slide list and rebuild in our order
    for child in list(sldIdLst):
        sldIdLst.remove(child)

    for rid, _ in DECK_ORDER:
        if rid in existing:
            sldIdLst.append(existing[rid])

    # Write back
    new_xml = ET.tostring(root, encoding='unicode', xml_declaration=False)
    new_xml = "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + new_xml
    with open(prs_path, 'w', encoding='utf-8') as f:
        f.write(new_xml)

def build_pptx(sections, meta):
    """Build a PPTX from proposal sections and metadata. Returns path to temp file."""
    accent = detect_colour(meta.get('client', ''))
    venue = meta.get('venue', 'Project')
    contact = meta.get('contact', '')
    role = meta.get('role', '')
    date_s = meta.get('date', '')

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError('Template not found — ensure 2020_template_slim_b64.txt is deployed.')

    with open(TEMPLATE_PATH, 'rb') as f:
        template_bytes = base64.b64decode(f.read())

    tmpdir = tempfile.mkdtemp(prefix='2020_')
    tpl_pptx = os.path.join(tmpdir, 'template.pptx')
    unpacked = os.path.join(tmpdir, 'unpacked')
    os.makedirs(unpacked)

    with open(tpl_pptx, 'wb') as f:
        f.write(template_bytes)

    with zipfile.ZipFile(tpl_pptx, 'r') as z:
        z.extractall(unpacked)
        slide14_raw = z.read('ppt/slides/slide14.xml').decode('utf-8')
        slide26_raw = z.read('ppt/slides/slide26.xml').decode('utf-8')

    # Convert from .potx template to .pptx presentation format
    # The template file has content type "presentationml.template" — PowerPoint
    # refuses to open it as a regular file. Patch both places that declare this.
    ct_path = os.path.join(unpacked, '[Content_Types].xml')
    with open(ct_path, encoding='utf-8') as f:
        ct = f.read()
    ct = ct.replace(
        'application/vnd.openxmlformats-officedocument.presentationml.template.main+xml',
        'application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml'
    )
    with open(ct_path, 'w', encoding='utf-8') as f:
        f.write(ct)

    # Also patch _rels/.rels — change template relationship type to presentation
    rels_path = os.path.join(unpacked, '_rels', '.rels')
    if os.path.exists(rels_path):
        with open(rels_path, encoding='utf-8') as f:
            rels = f.read()
        rels = rels.replace(
            '/relationships/presentationml/template',
            '/relationships/presentationml/presentation'
        )
        with open(rels_path, 'w', encoding='utf-8') as f:
            f.write(rels)

    # Fix viewProps.xml — template was saved in slide master view, change to normal view
    vp_path = os.path.join(unpacked, 'ppt', 'viewProps.xml')
    if os.path.exists(vp_path):
        with open(vp_path, encoding='utf-8') as f:
            vp = f.read()
        vp = vp.replace('lastView="sldMasterView"', 'lastView="sldView"')
        with open(vp_path, 'w', encoding='utf-8') as f:
            f.write(vp)

    # Reorder slides — keep only our 10 proposal slides in the right order
    reorder_presentation(unpacked)

    # Replace accent2 in all themes
    themes_dir = os.path.join(unpacked, 'ppt', 'theme')
    if os.path.exists(themes_dir):
        for tf in os.listdir(themes_dir):
            if tf.endswith('.xml'):
                fpath = os.path.join(themes_dir, tf)
                with open(fpath) as f:
                    tx = f.read()
                tx = re.sub(
                    r'<a:accent2><a:srgbClr val="[0-9A-Fa-f]{6}"/></a:accent2>',
                    f'<a:accent2><a:srgbClr val="{accent}"/></a:accent2>',
                    tx
                )
                tx = replace_colour(tx, 'E97132', accent)
                if tx.startswith('<?xml') and 'unicode' in tx[:60]:
                    tx = tx[tx.index('?>')+2:].lstrip()
                    tx = "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + tx
                with open(fpath, 'w', encoding='utf-8') as f:
                    f.write(tx)

    def load_slide(n):
        with open(os.path.join(unpacked, 'ppt', 'slides', f'slide{n}.xml')) as f:
            return ET.fromstring(f.read())

    def save_slide(n, root_elem):
        # Strip ET's invalid 'encoding=unicode' declaration and use correct UTF-8 one
        raw = ET.tostring(root_elem, encoding='unicode')
        if raw.startswith('<?xml'):
            raw = raw[raw.index('?>')+2:].lstrip()
        xml = "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + raw
        xml = replace_colour(xml, 'E8251A', accent)
        with open(os.path.join(unpacked, 'ppt', 'slides', f'slide{n}.xml'), 'w', encoding='utf-8') as f:
            f.write(xml)

    def save_slide_str(n, xml_str):
        # Fix declaration if present
        if xml_str.startswith('<?xml'):
            xml_str = xml_str[xml_str.index('?>')+2:].lstrip()
        xml_str = "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + xml_str
        xml_str = replace_colour(xml_str, 'E8251A', accent)
        with open(os.path.join(unpacked, 'ppt', 'slides', f'slide{n}.xml'), 'w', encoding='utf-8') as f:
            f.write(xml_str)

    # Build slides
    save_slide(1, build_cover(load_slide(1), venue, contact, role, date_s))
    save_slide(27, build_hello(load_slide(27)))

    brief = find_section(sections, 'your brief', 'brief reflection', 'understanding')
    save_slide(8, build_content_slide(
        load_slide(8), 'Your brief', 'Our understanding',
        first_sentences(brief, 2), explicit_bullets(brief, 5) or []
    ))

    save_slide(28, build_dark_divider(load_slide(28), 'Our methodology'))

    for slide_n, key1, key2, label in [
        (14, 'stage 1', 'strategic framework', 'Stage 1 — Strategic framework'),
        (9,  'stage 2', 'concept design',      'Stage 2 — Concept design'),
        (10, 'stage 3', 'design development',  'Stage 3 — Design development'),
        (11, 'stages 4', 'design intent',       'Stages 4, 5 and 6'),
    ]:
        txt = find_section(sections, key1, key2)
        save_slide_str(slide_n, build_stage_slide(
            slide14_raw, 'Our methodology', label, txt, explicit_bullets(txt, 8), accent
        ))

    fees_stages = [
        {'title': 'STRATEGIC FRAMEWORK', 'sub': 'Workshop, site visit and proposition'},
        {'title': 'CONCEPT DESIGN',      'sub': 'Layouts, materials, CGI visuals'},
        {'title': 'DESIGN DEVELOPMENT',  'sub': 'Sample boards and concept freeze'},
        {'title': 'DESIGN INTENT & ARTWORK', 'sub': 'Drawing pack and graphic artwork'},
    ]
    save_slide(16, build_fees_slide(load_slide(16), fees_stages, accent))
    save_slide_str(26, replace_colour(slide26_raw, 'E8251A', accent))

    # Pack into PPTX — [Content_Types].xml must be first, then _rels/.rels
    # Use the original template as base and patch modified slides in
    output_path = os.path.join(tmpdir, 'output.pptx')

    # Build a map of modified files
    modified = {}
    for root_dir, dirs, files in os.walk(unpacked):
        for file in files:
            fp = os.path.join(root_dir, file)
            arc = os.path.relpath(fp, unpacked).replace(os.sep, '/')
            with open(fp, 'rb') as f:
                modified[arc] = f.read()

    # Write zip with correct ordering: content types first, then rels, then everything else
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        # 1. Content types must be first
        if '[Content_Types].xml' in modified:
            zout.writestr('[Content_Types].xml', modified['[Content_Types].xml'])
        # 2. Root rels
        if '_rels/.rels' in modified:
            zout.writestr('_rels/.rels', modified['_rels/.rels'])
        # 3. Everything else
        for arc, data in modified.items():
            if arc not in ('[Content_Types].xml', '_rels/.rels'):
                zout.writestr(arc, data)

    return output_path, tmpdir


# ── AI PIPELINE ───────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are a proposal writer for 20.20 Design Agency, a design consultancy specialising in hospitality interiors, stadium design, brand identity, graphics, environmental design, and wayfinding. You write first-draft proposals that the account team will review and refine.

The PROJECT TYPE field in context tells you what kind of proposal this is. Follow it precisely:
- HOSPITALITY / INTERIOR DESIGN: spatial language, guest journey, hospitality tiers, CGI renders, RIBA stages if applicable
- GRAPHICS AND BRAND: graphic design language, print specifications, artwork delivery, brand guidelines — no spatial plans, no CGI, no RIBA stages
- BRAND STRATEGY: strategic frameworks, workshop outputs, positioning documents — no physical deliverables
- FORMAL TENDER / ITT RESPONSE: formal tone, respond to evaluation criteria, include assumptions and exclusions
- CRUISE / FIT-OUT: shipyard phase gate language, technical specifications, information release schedule

If the project type is unclear or flagged as catering/operations, write conservatively and flag the uncertainty.

VOICE: Confident, direct, commercially aware, personal. Short sentences. Active voice. No em dashes. No AI phrases (leveraging, seamless, holistic, transformative). The client name appears only in the cover letter — all other sections refer to the client as "the client", "the organisation", or use the venue/project name where relevant. Never use "the club" for non-sports clients.

DESIGN PRINCIPLES: Hospitality Pyramid (tier each space), Narrative Before Design (names and stories before materials), Guest Journey Mapping, Brand Integration Without Decoration, Non-Matchday Flexibility, Graphic Identity as Interior Design, Commercial Consciousness, CGI from Stage 2, Collaborative Design Team, Concept Freeze.

PLACEHOLDERS: [FEE: TBC — rate card required] for fees, [IMAGE REQUIRED: description] for visuals, [CONFIRM WITH CLIENT: note] for assumptions."""

# Stage prompt template — generates structured content with four named sections
# Works for both RIBA-staged and phase-based proposals
STAGE_PROMPT = """Write the {stage_name} section for this 20.20 proposal.

THIS IS CRITICAL: Write specifically to THIS brief, not generically about hospitality design.
Reference the actual spaces listed. Reference the actual tiers (Bronze, Silver, Gold, VVIP).
Reference the specific brief requirements, constraints and programme.
If there is a lead concept space model or specific design approach requested, reflect that.
If certain things are fixed (seat positions, lounge sizes, kitchen layouts), acknowledge those constraints.

Structure your response with exactly these four labelled sections:

Objective:
[1-2 sentences on what this stage achieves for THIS project — reference the specific venue, spaces or tiers]

Process:
[4-6 bullet points — how we work through this stage. Reference the specific spaces, tiers and design team context where relevant. If the project has multiple spaces or sub-stages (e.g. 3.1, 3.2, 3.3), summarise the sequence and grouping logic here in a few bullets — do not give each sub-stage its own bullet if that means Deliverables ends up empty.]

Deliverables:
[6-10 bullet points — specific outputs, covering ALL spaces and sub-stages named in this stage, not just the ones mentioned in Process. If there are named spaces, reference them. If there is a per-tier delivery model, reflect it. Include the number of CGI renders if concept stage]

Meetings & Presentations:
[3-5 bullet points — specific meetings with the design team, architect and client. Reference Teams or in-person based on what the brief says]

CRITICAL — sub-stages and multi-space projects: when a stage covers multiple spaces or numbered sub-stages (e.g. Stage 3.1 through 3.6 across different lounges), do NOT spread the sub-stage walkthrough across Objective, Process and Deliverables as if each column holds a different slice of the sequence. Each of Objective, Process, Deliverables and Meetings must independently cover the FULL stage, every space and every sub-stage, just from that column's own angle (what it achieves / how we do it / what we produce / who we meet). A reader looking at only the Deliverables column should see outputs for every single space in this stage, not just some of them.

Quality requirements — these are mandatory:

1. RESPOND TO ALL KEY REQUIREMENTS: Every requirement and deliverable the client stated in the brief must be addressed in this stage section. Do not omit anything they asked for. If the brief specifies a particular output (e.g. fly-throughs, a specific report format, a named deliverable), it must appear in the Deliverables list.

2. PROPORTIONATE RESPONSE: Give more detail on the things the brief emphasises. If the client spent a paragraph on bar positioning, address it properly. If they flagged a specific constraint, acknowledge it. Weight your response to match the importance placed on topics in the brief.

3. USE THE CLIENT'S TERMINOLOGY: Use the exact words and phrases from the brief. If they call it "the James Herriot Restaurant", use that name — not "the restaurant". If they say "Bronze, Silver, Gold", use those exact tier names. If they refer to "the Demand and Revenue Assessment", use that phrase. Mirror their technical and factual language precisely. Then write in 20.20 tone of voice around it — confident, direct, commercially aware.

4. TIMINGS: Use realistic design timings. Stage 1 / Phase 1 = 2-3 weeks. Stage 2 / Phase 2 = 4-6 weeks. Stage 3 / Phase 3 = 6-8 weeks. Stage 4 / Phase 4 = 8-12 weeks. Stages 5-6 = programme dependent. Do not use 1-2 week timings — these are too short.

Format rules:
- No markdown. No asterisks. No bold (**text**). No headers (#).
- Write the four section labels as plain text on their own line followed by a colon.
- The Meetings section must name specific meeting types and cadence — not a placeholder.
- Use "the client", "the organisation", "the venue", or "the project" — not the client name, and never "the club" for non-sports clients.

{ctx}"""


SECTIONS = [
            ('cover',   'Cover letter',
     '''Write the cover letter for this 20.20 proposal. It appears as the "Hello" slide.

First line must be: Hello [first name of contact] (no full stop, no "Dear")
If two contacts: Hello [Name] and [Name]

Then: Dear [Name], (on a new line, this is the actual letter salutation)

Write 3-4 SHORT paragraphs. Maximum 3 sentences each. No filler.

If this is a CONTINUATION or REBRIEF (continuation=yes or prior stages noted):
  Para 1: Warmly acknowledge the prior relationship and work done. Be specific if prior stages are known.
  Para 2: What this proposal covers — the specific stages/phases, the specific spaces or areas.
  Para 3: What we need from the client to deliver it (clear decisions, access, timescale).
  Para 4 (optional): Brief confident close.

If this is a NEW brief:
  Para 1: What drew 20.20 to this project — something specific about the venue or brief.
  Para 2: What this proposal covers.
  Para 3: What we need from the client.

Sign off: "Kind regards," on its own line, then "The 20.20 team" on the next line.

Rules:
- No markdown. No asterisks. No bullet points.
- Short sentences. No em dashes. Direct tone.
- The client name or venue name appears at most once.
- Do NOT write generic agency positioning paragraphs about 20.20's methodology.
- Use the specific language from the brief — if they named a programme or a deadline, reference it.
- Timings must be realistic: Stage 2 = 4-6 weeks, Stage 3 = 6-8 weeks — not 1-2 weeks.

{ctx}'''),
    ('brief',   'Your brief',
     '''Write the "Your brief" section. This shows the client we listened and understood exactly what they said.

This section must be SPECIFIC and VERBATIM where possible. Do NOT write a strategic analysis.

Structure:
- Para 1 (2-3 sentences): The overall project in plain commercial terms. Include any specific numbers from the brief (capacity, revenue targets, budget if stated).
- Then one short paragraph per KEY SPACE or AREA, in order of importance. For each:
  * Name the space exactly as the brief names it
  * State its tier/level in the hospitality hierarchy
  * State its capacity or commercial purpose
  * Note any specific requirement, constraint or client preference mentioned

If the brief names constraints (fixed lounge sizes, seat positions, kitchen areas that cannot move), state them.
If the client mentioned dislikes or things to avoid, include them.
If spaces are not named in the brief, acknowledge this and note what IS known.

Use "the client", "the organisation", or "the venue" — not the client name, and never "the club" for non-sports clients.
No markdown. No bold text. No bullet points — write in short paragraphs.
Use the EXACT TERMINOLOGY from the brief — if they name a space, use that exact name. If they give capacity figures or budget ranges, include them verbatim.
Flag anything that needs confirming with: [CONFIRM WITH CLIENT: what needs clarifying]

{ctx}'''),
    ('stage1',  'Stage 1 — Strategic framework',
     STAGE_PROMPT.format(stage_name='Stage 1 — Strategic framework (RIBA Stage 2, 1 week)', ctx='{ctx}')),
    ('stage2',  'Stage 2 — Concept design',
     STAGE_PROMPT.format(stage_name='Stage 2 — Concept design (RIBA Stage 2, 2 weeks). Include CGI commitment — minimum 2 visuals per space in Deliverables', ctx='{ctx}')),
    ('stage3',  'Stage 3 — Design development',
     STAGE_PROMPT.format(stage_name='Stage 3 — Design development (RIBA Stage 2, 2 weeks). Include concept freeze milestone in Deliverables', ctx='{ctx}')),
    ('stage456','Stage 4 onwards',
     STAGE_PROMPT.format(stage_name='The final stage(s) of the project covering technical production, coordination and handover. For RIBA-staged projects call this "Stages 4, 5 and 6" and use sub-headings Stage 4 / Stage 5 / Stage 6. For phase-based or arena projects call this "Phase 4 — Production and delivery" and describe it as a single phase. Match the naming convention used in the earlier stage sections. IMPORTANT: this section covers multiple sub-stages in one response. Every sub-stage MUST have its own populated Deliverables list — do not let Objective and Process content crowd out Deliverables. If space is tight, keep Process bullets shorter rather than dropping Deliverables content.', ctx='{ctx}')),
    ('fees',    'Fees and timings',
     'Write the fees section. List each stage with [FEE: TBC] for all figures. Note timings per stage. Fees are exclusive of VAT, 3rd party costs, general expenses and travel. Subject to contract.\n\n{ctx}'),
    ('nextsteps','Next steps',
     'Write next steps as four numbered actions: review, feedback, site visit, appointment. '
     'Format each one as a short title followed by a colon, then a 1-2 sentence description, '
     'for example: "Review this proposal: Share it with your team and flag any questions before we proceed." '
     'Keep the title under 5 words. Direct and confident. No client name.\n\n{ctx}'),
]

# ── GRAPHICS / BRAND TEMPLATE ─────────────────────────────────────────────────
# Used when project_type is 'graphics_brand' or 'other' (non-hospitality).
# Shorter, fewer stages, no spatial language, no CGI, no RIBA references.
# Deliverables are artwork files, guidelines, print-ready assets.
SECTIONS_GRAPHICS = [
    ('cover', 'Cover letter',
     '''Write the cover letter for this 20.20 proposal. It appears as the "Hello" slide.

First line: Hello [first name of contact]
Then: Dear [Name],

Write 3 SHORT paragraphs. Maximum 3 sentences each.

Para 1: What drew 20.20 to this project — something specific about the brief or the brand challenge.
Para 2: What this proposal covers — the specific deliverables (be concrete: gift cards, brand guidelines, artwork files etc.)
Para 3: What we need from the client to deliver it (VBI/brand guidelines access, copy approval, timeline, sign-off process).

Sign off: "Kind regards," then "The 20.20 team"

Rules:
- No markdown. No asterisks. No bullet points.
- Short sentences. No em dashes. Direct tone.
- Never use "the club" — use "the brand", "the organisation" or the client name in the first mention only.
- Do NOT reference spatial design, interior design, RIBA stages, CGI, guest journeys or stadium language.
- This is a graphics and brand project. Write accordingly.

{ctx}'''),

    ('brief', 'Your brief',
     '''Write the "Our understanding" section for this graphics/brand proposal.

Two columns of prose (no bullets in left column, bullets allowed in right column).

LEFT COLUMN — Summarise the brief in plain English:
- What the client wants made (be specific: gift cards, brand identity, artwork, guidelines etc.)
- The brand context — what repositioning, campaign or strategic moment this sits within
- The emotional or commercial brief — what the work needs to achieve beyond just looking good
- Any constraints named (brand guidelines to follow, technical specs, approvals process)

RIGHT COLUMN — Space-by-space or piece-by-piece breakdown:
- For each deliverable (card, sleeve, format, variant) give a brief note on its specific requirements
- Flag anything that needs confirming with [CONFIRM WITH CLIENT: specific question]

Rules:
- No spatial language. No "guest journey", no "lounge", no "tier", no "hospitality".
- This is a graphics and brand project. Use brand, design, artwork, print language throughout.
- Never use "the club" — use "the brand" or "the organisation".
- Use the exact names the brief uses for deliverables.

{ctx}'''),

    ('stage1', 'Phase 1 — Concept design',
     '''Write Phase 1 — Concept design for this graphics/brand proposal.

This phase covers initial creative direction: reviewing the brief and brand guidelines, developing
concept routes, presenting options and getting a single direction signed off before any artwork begins.

Structure with exactly these four labelled sections:

Objective:
[1-2 sentences: what Phase 1 achieves — a signed-off creative direction before production begins]

Process:
- Brief audit and brand review: review all brand guidelines, VBI, existing assets and any technical specs
- Concept development: develop [2-3] distinct concept routes covering visual language, typography, colour approach and hierarchy
- Internal review before presenting to client
- Client presentation: present concept routes, gather feedback, agree a single direction
- Direction sign-off: confirm the chosen direction in writing before Phase 2 begins

Deliverables:
- [number] concept routes presented as visual boards showing design language, typography, colour and tone
- Written rationale for each route explaining how it meets the brief
- Agreed and signed-off creative direction document
- Any brand compliance notes flagged before production begins

Meetings & Presentations:
- Kick-off call at start of phase
- Concept presentation to client team
- Sign-off confirmation before Phase 2 commences
- All documents issued via PDF

Format rules: No markdown. No asterisks. No RIBA language. No spatial language. No CGI.
Use "the brand" or "the organisation" not "the club".

{ctx}'''),

    ('stage456', 'Phase 2 — Artwork and delivery',
     '''Write Phase 2 — Artwork and delivery for this graphics/brand proposal.

This phase covers production of the final artwork: taking the approved concept through to
print-ready or screen-ready files, including all technical checks, colour proofing and client sign-off.

Structure with exactly these four labelled sections:

Objective:
[1-2 sentences: what Phase 2 achieves — client-approved, production-ready artwork files]

Process:
- Artwork preparation: produce all components simultaneously to maintain visual consistency
- Work to confirmed physical dimensions, substrate specs or screen specs as required
- Colour management: ensure colour palette holds correctly across all components and materials
- Copy and content lock: confirm all copy, legal lines and variable data in writing before final file build
- Pre-press review or technical check on all components before submission
- File packaging and handover in agreed formats

Deliverables:
- Final artwork files for all components (list each one specifically from the brief)
- Colour proofs for client written sign-off before production release
- Copy-approved and locked text document
- Pre-press checklist signed off internally
- Final packaged files supplied in agreed format (AI, PDF, CMYK as required)
- Print specification summary or technical handover document

Meetings & Presentations:
- Mid-phase review at artwork stage
- Sign-off on colour proofs before release
- File handover and production briefing
- All documents issued via PDF

Format rules: No markdown. No asterisks. No RIBA language. No spatial language. No CGI, no sample boards.
Use print and artwork language throughout. Use "the brand" or "the organisation" not "the club".

{ctx}'''),

    ('fees', 'Fees and timings',
     '''Write the fees and timings summary for this graphics/brand proposal.

List the phases with realistic timings for a graphics/design project:
- Phase 1 — Concept design: 2-3 weeks
- Phase 2 — Artwork and delivery: 2-3 weeks

Fees are [FEE: TBC] for all phases unless stated in the brief.
Invoicing: 50% at start of each phase, 50% on completion.
Note: fees are exclusive of VAT, third-party costs and print production.

No markdown. No RIBA language.

{ctx}'''),

    ('nextsteps', 'Next steps',
     '''Write next steps as four numbered actions for this graphics/brand proposal.
Format each as: "Title: one sentence description."
1. Review this proposal: ...
2. Send us your feedback: ...
3. Share brand assets: ask them to share brand guidelines, VBI or any existing assets we need before starting
4. Confirm appointment: ...

Keep titles under 5 words. Direct and confident. No client name. No spatial language.

{ctx}'''),
]


def build_context(meta, spaces_text=''):
    bt = meta.get('brief_type','')
    riba = meta.get('riba_stages','')
    continuation = meta.get('continuation','no')
    prior = meta.get('prior_stages_completed','')
    second = meta.get('second_contact','')
    project_type = meta.get('project_type', '')

    # Stage context line
    if riba:
        stage_ctx = f"RIBA STAGES REQUESTED: {riba}"
        if meta.get('stage_2_duration'): stage_ctx += f" | Stage 2: {meta['stage_2_duration']}"
        if meta.get('stage_3_duration'): stage_ctx += f" | Stage 3: {meta['stage_3_duration']}"
    else:
        stage_ctx = f"STAGES/PHASES: {riba or 'To be confirmed'}"

    contact_line = meta.get('contact','')
    if meta.get('role'): contact_line += f", {meta['role']}"
    if second: contact_line += f" and {second}"

    is_riba_flag = meta.get('is_riba','yes').lower()
    ctx = f"PROJECT: {meta.get('venue','')}\n"

    # Project type — the primary signal for template and terminology selection.
    # Written explicitly so every section prompt knows what kind of proposal this is.
    pt_labels = {
        'hospitality':       'HOSPITALITY / INTERIOR DESIGN — use spatial language, guest journey, tiers, hospitality pyramid',
        'graphics_brand':    'GRAPHICS AND BRAND — use graphic design language. No CGI, no spatial plans, no RIBA stages. Deliverables are artwork files, guidelines, print-ready assets.',
        'strategy_brand':    'BRAND STRATEGY — use strategy and brand language. Deliverables are frameworks, guidelines, workshop outputs. No physical design deliverables.',
        'tender_itt':        'FORMAL TENDER / ITT RESPONSE — formal tone, respond to stated selection criteria, include assumptions and exclusions, structure around the client\'s evaluation framework.',
        'cruise_fitout':     'CRUISE / FIT-OUT — structure around shipyard phase gates and information release schedule. No hospitality pyramid. Deliverables are technical drawings and specifications.',
        'catering_operations': 'CATERING / OPERATIONS — NOTE: this does not appear to be a design brief. Confirm the actual scope before generating.',
        'unknown':           'PROJECT TYPE UNCLEAR — write with caution, flag assumptions.',
    }
    pt_label = pt_labels.get(project_type, f'PROJECT TYPE: {project_type or "not determined"}')
    ctx += f"PROJECT TYPE: {pt_label}\n"

    ctx += f"RIBA STAGED PROJECT: {'YES — respond using RIBA stage structure and terminology' if is_riba_flag == 'yes' else 'NO — this is a phase-based or single-scope project, do not use RIBA stage references'}\n"
    ctx += f"CLIENT: {meta.get('client','')}\n"
    if meta.get('sector'):
        ctx += f"CLIENT SECTOR: {meta['sector']} — use language and terminology appropriate to this sector throughout. Never use sports club language for non-sports clients.\n"
    ctx += f"CONTACT: {contact_line}\n"
    if meta.get('lead_architect'): ctx += f"LEAD ARCHITECT/DESIGN TEAM: {meta['lead_architect']}\n"
    if meta.get('project_manager'): ctx += f"PROJECT MANAGER: {meta['project_manager']}\n"
    ctx += f"BRIEF TYPE: {bt}\n"
    ctx += f"BRIEF SOURCE: {meta.get('brief_source','')}\n"
    ctx += f"CONTINUATION OF PRIOR WORK: {continuation.upper()}\n"
    if prior: ctx += f"PRIOR STAGES / CONTEXT: {prior}\n"
    ctx += f"{stage_ctx}\n"
    ctx += f"BUDGET: {meta.get('budget','Not stated')}\n"
    if meta.get('tier_summary'): ctx += f"HOSPITALITY TIERS: {meta['tier_summary']}\n"

    # Graphics/brand specific context
    if meta.get('deliverable_types'):
        ctx += f"DELIVERABLES EXPECTED: {meta['deliverable_types']}\n"
    if meta.get('print_required') == 'yes':
        ctx += "PRINT REQUIRED: Yes — include CMYK, print-ready artwork, file format specs in deliverables.\n"
    if meta.get('brand_guidelines_exist') == 'yes':
        ctx += "EXISTING BRAND GUIDELINES: Yes — design must work within confirmed brand guidelines, not develop new ones from scratch.\n"
    elif meta.get('brand_guidelines_exist') == 'no':
        ctx += "EXISTING BRAND GUIDELINES: No — brand language and visual system to be developed as part of this project.\n"

    # Internal notes warning — so the model doesn't repeat internal content in the proposal
    if meta.get('contains_internal_notes') == 'yes':
        ctx += (f"\nINTERNAL NOTES DETECTED: The brief appears to contain internal observations "
                f"that should NOT appear in the client-facing proposal. "
                f"{meta.get('internal_notes_description','')} "
                f"Focus only on the actual client requirements and project scope.\n")

    # Spaces
    if spaces_text:
        ctx += f"SPACES:\n{spaces_text}\n"
    elif meta.get('scope_plain'):
        ctx += f"SCOPE: {meta['scope_plain']}\n"
    elif meta.get('scope'):
        ctx += f"SCOPE: {meta['scope']}\n"

    # Key brief content
    if meta.get('key_requirements'):
        ctx += f"\nKEY REQUIREMENTS FROM THE BRIEF:\n{meta['key_requirements']}\n"
    if meta.get('key_constraints'):
        ctx += f"\nCONSTRAINTS (fixed elements, things NOT to change):\n{meta['key_constraints']}\n"
    if meta.get('client_dislikes'):
        ctx += f"\nCLIENT DISLIKES/AVOIDED APPROACHES:\n{meta['client_dislikes']}\n"
    if meta.get('design_approach'):
        ctx += f"\nSPECIFIC DESIGN APPROACH REQUESTED:\n{meta['design_approach']}\n"
    if meta.get('supporting_context'):
        ctx += f"\nADDITIONAL CONTEXT FROM SUPPORTING DOCUMENTS (the primary brief above always takes priority if anything conflicts):\n{meta['supporting_context']}\n"

    return ctx.strip()


def repair_json(raw_str):
    """Try progressively more aggressive repairs on near-valid JSON from the model.
    Handles the most common failure: unescaped quotes/apostrophes inside string
    values (club names, quoted brief language), which break json.loads with
    'Expecting , delimiter' type errors."""
    # Attempt 1: as-is
    try:
        return json.loads(raw_str)
    except json.JSONDecodeError:
        pass

    # Attempt 2: escape stray backslashes that aren't valid escape sequences
    try:
        fixed = re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', raw_str)
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    # Attempt 3: fix smart quotes / curly apostrophes inside values, which
    # sometimes appear when the model echoes brief text verbatim
    try:
        fixed = raw_str.replace('\u2018', "'").replace('\u2019', "'")
        fixed = fixed.replace('\u201c', '"').replace('\u201d', '"')
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    # Attempt 4: line-by-line repair — find the specific line/column json
    # reports and escape an unescaped " or ' that isn't a field delimiter.
    try:
        fixed = raw_str
        for _ in range(8):  # cap repair attempts to avoid infinite loop
            try:
                return json.loads(fixed)
            except json.JSONDecodeError as e:
                idx = e.pos
                # If the character at the error position is an unescaped quote
                # inside what looks like a string value, escape it and retry.
                if idx < len(fixed) and fixed[idx] == '"':
                    fixed = fixed[:idx] + '\\"' + fixed[idx+1:]
                else:
                    raise
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    return None


def strip_html(txt):
    """Strip HTML tags from text — intel comes back with <cite> tags from web search."""
    if not txt:
        return ''
    import re as _re
    clean = _re.sub(r'<[^>]+>', '', str(txt))
    clean = clean.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&nbsp;', ' ')
    return ' '.join(clean.split())

def run_pipeline(job_id, pdf_b64=None, brief_text=None, prior_work='', supporting_docs_b64=None):
    """Background thread: extract → research → generate → build PPTX."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
    supporting_docs_b64 = supporting_docs_b64 or []
    pipeline_start_ts = time.time()

    log_usage_event(
        job_id, 'started',
        input_type='pdf' if pdf_b64 else 'text',
        num_supporting_docs=len(supporting_docs_b64),
        has_prior_work=bool(prior_work),
    )

    def progress(msg, pct=None):
        append_progress(job_id, msg, pct)

    try:
        # ── STEP 1: EXTRACT MAIN BRIEF ───────────────────────────────────────
        # Supporting docs are deliberately NOT stacked into this call. Each
        # extra PDF in the same request multiplies the input tokens for that
        # single call, and on lower usage tiers (30k input tokens/minute is
        # the default tier) a brief plus 2-3 supporting PDFs can blow past
        # the limit in one shot, where retrying just fails again at the same
        # size. Instead: extract the main brief alone (small, fast, well
        # within limits), then summarize each supporting doc in its OWN
        # separate, smaller call below, spacing token usage out over time
        # rather than concentrating it in one request.
        progress('Reading the brief...', 5)
        extract_prompt = (
            'Read this client brief, ITT, scope document, or meeting notes carefully. '
            'Extract ALL available information. Return ONLY valid JSON with NO markdown or explanation.\n'
            'CRITICAL: every string value must be valid JSON. Escape all double quotes as \\" and avoid '
            'using straight or curly apostrophes inside values where possible. Do not include literal '
            'line breaks inside a string value — use a single space instead.\n\n'

            'STEP 1 — DETERMINE PROJECT TYPE FIRST. This is the most important decision. '
            'Look for these vocabulary signals before reading anything else:\n'
            '  HOSPITALITY/INTERIOR: lounge, bar, restaurant, concourse, tier, hospitality, RIBA stage, '
            'CGI, spatial, interior design, FF&E, furniture, materials, guest journey, seating bowl\n'
            '  GRAPHICS/BRAND: gift card, packaging, sleeve, artwork, print, CMYK, brand identity, '
            'tone of voice, guidelines, logo, visual system, typographic, colour palette, brand strategy, '
            'positioning, propositions, concept routes, templates, signage design, wayfinding\n'
            '  STRATEGY/BRAND ONLY (no physical deliverables): workshop, archetypes, key messages, '
            'competitive intelligence, market positioning — with NO spatial or print deliverables mentioned\n'
            '  TENDER/ITT: ITT, ITN, RFP, tender submission, selection criteria, evaluation criteria, '
            'Form of Tender, fee return, submission deadline\n'
            '  CATERING/OPERATIONS (not a design brief): catering contract, menu, kitchen operator, '
            'F&B provision, catering tender — flag this as catering_operations\n'
            '  CRUISE/FIT-OUT: shipyard, yard phase, hull number, Engineering Unit, Meyer Werft, '
            'fit-out, cabin typology, GAP drawing, wayfinding on a vessel\n\n'

            'STEP 2 — ASSESS CONFIDENCE on critical fields. For each, set confidence to:\n'
            '  high: clear, unambiguous, stated explicitly\n'
            '  medium: inferable but not explicit, or one strong candidate among minor alternatives\n'
            '  low: genuinely ambiguous, multiple equally plausible answers, or absent entirely\n\n'

            'STEP 3 — DETECT INTERNAL NOTES. Meeting notes, internal strategy discussions, budget '
            'observations, staff changes, competitor intelligence, and internal opinions should NOT '
            'appear in a client-facing proposal. Flag if present.\n\n'

            '{\n'

            # ── PROJECT TYPE (new unified field) ───────────────────────────
            '"project_type": "hospitality | graphics_brand | strategy_brand | tender_ itt | '
            'cruise_fitout | catering_operations | unknown — use the vocabulary signals above. '
            'hospitality covers all interior design work regardless of project scale. '
            'graphics_brand covers print, packaging, signage design, environmental graphics, brand identity. '
            'strategy_brand is brand strategy with no physical deliverables. '
            'tender_itt is any formal tender or ITT response regardless of discipline. '
            'catering_operations means this is not a design brief at all.",\n'
            '"project_type_confidence": "high | medium | low",\n'
            '"project_type_question": "if confidence is medium or low, write the exact plain-language '
            'question to show the user — e.g. This could be a hospitality or graphics project — which is '
            'right? Leave empty string if confidence is high.",\n\n'

            # ── EXISTING FIELDS (kept, not broken) ─────────────────────────
            '"is_riba": "yes or no. yes = RIBA Plan of Work stages explicitly referenced. '
            'no = phase-based, single-scope, brand, graphics, or no RIBA stage references.",\n'
            '"brief_type": "newbuild | refurb | single_space | sponsor | arena | continuation | '
            'itt | graphics | brand | cruise | unknown",\n'
            '"brief_source": "Direct approach | Via architect or PM | Formal open tender (ITT) | '
            'Referral | Repeat client | Unknown",\n'
            '"continuation": "yes | no",\n'

            # ── CLIENT (with confidence) ────────────────────────────────────
            '"client": "the name a 20.20 proposal would be addressed to — usually the organisation '
            'commissioning the work. If meeting notes mention both a venue operator and a tenant or '
            'sponsor, choose the one doing the commissioning. If both are plausible, set confidence low.",\n'
            '"client_confidence": "high | medium | low",\n'
            '"client_question": "if confidence is medium or low, plain-language question — e.g. The '
            'brief mentions both Mattioli Woods and Leicester Tigers — which should we use as the client '
            'name? Leave empty string if confidence is high.",\n\n'
            '"sector": "the industry or world the client operates in — a plain user-friendly label. '
            'Examples: Professional football club, Coffee brand and retailer, Event catering company, '
            'Cruise line, Property developer, Hospitality venue operator, Arena operator, Brand agency client. '
            'NEVER use the word club for non-sports organisations. '
            'NEVER default to sports or stadium language for brands, retailers or other sectors.",\n\n'

            '"venue": "",\n'
            '"primary_contact": "first name only if clear, otherwise full name — used in Dear [Name] greeting",\n'
            '"contact_role": "",\n'
            '"second_contact": "",\n'
            '"lead_architect": "name of lead architect or design team lead if mentioned",\n'
            '"project_manager": "name of PM or client representative if mentioned",\n'
            '"proposal_deadline": "",\n'
            '"construction_completion": "",\n'
            '"budget_stated": "",\n'
            '"riba_stages": "exact RIBA stages requested e.g. Stage 2 and 3 — be precise. '
            'Leave empty if not a RIBA project.",\n'
            '"stage_2_duration": "weeks if stated",\n'
            '"stage_3_duration": "weeks if stated",\n'
            '"prior_stages_completed": "any stages already done e.g. Stage 2 complete",\n'
            '"spaces": [{"name": "", "tier": "Bronze|Silver|Gold|VVIP|GA|GA+", "level": "", '
            '"capacity": "", "budget": "", "notes": "specific requirements for this space"}],\n'
            '"tier_summary": "e.g. Gold 1372 seats, Silver 642. Leave empty if not a tiered hospitality project.",\n'
            '"key_requirements": "the actual client requirements and deliverables from the brief — '
            'NOT internal notes or 20.20 opinions. Verbatim or near-verbatim where possible.",\n'
            '"key_constraints": "fixed elements, things not to change, operational constraints",\n'
            '"client_dislikes": "anything client has said they do not want",\n'
            '"design_approach": "any specific approach the brief requests",\n'

            # ── GRAPHICS/BRAND SPECIFIC ─────────────────────────────────────
            '"deliverable_types": "list the actual deliverable types mentioned — e.g. gift card design, '
            'brand guidelines, print artwork, wayfinding drawings, CGI renders, signage package. '
            'Leave empty if not determinable.",\n'
            '"print_required": "yes | no | unknown — whether print-ready or CMYK artwork is needed",\n'
            '"brand_guidelines_exist": "yes | no | unknown — whether the client has existing brand '
            'guidelines 20.20 must work within",\n'

            # ── INTERNAL NOTES DETECTION ────────────────────────────────────
            '"contains_internal_notes": "yes | no — yes if the brief contains meeting notes, '
            'internal opinions, competitor intelligence, staff observations, or budget context '
            'that should NOT appear in a client-facing proposal",\n'
            '"internal_notes_description": "if contains_internal_notes is yes, one sentence '
            'describing what looks internal — e.g. Notes include staff departure details and '
            'internal budget observations. Leave empty string if no internal notes detected.",\n'

            # ── TRIAGE CONTROL FIELDS ───────────────────────────────────────
            '"scope_plain": "one plain-English sentence describing what work is in scope — '
            'written as a user would understand it, not in design jargon. E.g. Full four-phase '
            'hospitality design for the VVIP lounge at the Amex Stadium. Or: Brand strategy and '
            'identity development for Payne and Gunter across London event venues. Or: Design of '
            'two gift cards and sleeves for Costa Coffee thank you and birthday occasions.",\n'
            '"brief_summary": "one plain-English sentence the triage screen shows at the top — '
            'what the model understood at a glance. E.g. This looks like a four-phase hospitality '
            'proposal for Brighton and Hove Albion FC at the Amex Stadium. Or: This looks like a '
            'brand strategy project for Payne and Gunter — please confirm the project type before '
            'generating. Make the sentence more cautious if any confidence is medium or low.",\n'
            '"proceed_direct": "true | false — true only if project_type_confidence, '
            'client_confidence are BOTH high AND contains_internal_notes is no. '
            'false if ANY critical field is medium or low confidence, or internal notes detected. '
            'When false, the triage screen will show questions before generating."\n'
            '}'
        )

        if pdf_b64:
            msg_content = [
                {'type': 'document', 'source': {'type': 'base64', 'media_type': 'application/pdf', 'data': pdf_b64}},
                {'type': 'text', 'text': extract_prompt},
            ]
        else:
            msg_content = extract_prompt + '\n\nBrief:\n' + (brief_text or '')[:4000]

        resp = None
        for attempt in range(4):
            try:
                if attempt > 0:
                    wait = [0, 25, 45, 70][attempt]
                    progress(f'Rate limit reached — retrying extraction in {wait}s...', 5)
                    time.sleep(wait)
                resp = client.messages.create(
                    model='claude-sonnet-4-6',
                    max_tokens=1400,
                    messages=[{'role': 'user', 'content': msg_content}]
                )
                break
            except anthropic.RateLimitError:
                if attempt == 3:
                    raise ValueError(
                        'Rate limit reached reading the brief. Try again in a minute.'
                    )
        raw = resp.content[0].text.replace('```json', '').replace('```', '').strip()
        m = re.search(r'\{[\s\S]*\}', raw)
        if not m:
            raise ValueError('Could not extract brief data from the document.')
        ex = repair_json(m.group(0))
        if ex is None:
            # Last resort: ask the model to re-emit the same JSON, valid this time
            progress('Brief data needed a re-pass — retrying...', 8)
            fix_resp = client.messages.create(
                model='claude-sonnet-4-6',
                max_tokens=1400,
                messages=[{'role': 'user', 'content':
                    'The following text should be valid JSON but failed to parse. '
                    'Return ONLY the corrected, valid JSON with no markdown, no explanation, '
                    'and all string values properly escaped (especially apostrophes and quotes):\n\n'
                    + m.group(0)}]
            )
            fix_raw = fix_resp.content[0].text.replace('```json', '').replace('```', '').strip()
            fix_m = re.search(r'\{[\s\S]*\}', fix_raw)
            ex = repair_json(fix_m.group(0)) if fix_m else None
            if ex is None:
                raise ValueError('Could not parse brief data as JSON, even after a repair pass.')
        update_job(job_id, extracted=ex)

        meta = {
            'is_riba':               ex.get('is_riba', 'yes'),
            'client':                ex.get('client', ''),
            'sector':                ex.get('sector', ''),
            'venue':                 ex.get('venue', ''),
            'contact':               ex.get('primary_contact', ''),
            'role':                  ex.get('contact_role', ''),
            'second_contact':        ex.get('second_contact', ''),
            'lead_architect':        ex.get('lead_architect', ''),
            'project_manager':       ex.get('project_manager', ''),
            'brief_type':            ex.get('brief_type', ''),
            'brief_source':          ex.get('brief_source', ''),
            'continuation':          ex.get('continuation', 'no'),
            'prior_stages_completed':ex.get('prior_stages_completed', ''),
            'riba_stages':           ex.get('riba_stages', ''),
            'stage_2_duration':      ex.get('stage_2_duration', ''),
            'stage_3_duration':      ex.get('stage_3_duration', ''),
            'budget':                ex.get('budget_stated', ''),
            'tier_summary':          ex.get('tier_summary', ''),
            'scope':                 ex.get('scope_summary', ex.get('scope_plain', '')),
            'scope_plain':           ex.get('scope_plain', ''),
            'key_requirements':      ex.get('key_requirements', ''),
            'key_constraints':       ex.get('key_constraints', ''),
            'key_preferences':       ex.get('key_preferences', ex.get('key_requirements', '')),
            'client_dislikes':       ex.get('client_dislikes', ''),
            'design_approach':       ex.get('design_approach', ''),
            # ── New triage fields ───────────────────────────────────────────
            'project_type':          ex.get('project_type', ''),
            'project_type_confidence': ex.get('project_type_confidence', 'high'),
            'project_type_question': ex.get('project_type_question', ''),
            'client_confidence':     ex.get('client_confidence', 'high'),
            'client_question':       ex.get('client_question', ''),
            'contains_internal_notes': ex.get('contains_internal_notes', 'no'),
            'internal_notes_description': ex.get('internal_notes_description', ''),
            'deliverable_types':     ex.get('deliverable_types', ''),
            'print_required':        ex.get('print_required', 'unknown'),
            'brand_guidelines_exist':ex.get('brand_guidelines_exist', 'unknown'),
            'brief_summary':         ex.get('brief_summary', ''),
            'proceed_direct':        str(ex.get('proceed_direct', 'true')).lower() == 'true',
            'date':                  time.strftime('%-d %B %Y'),
        }

        # ── STEP 1b: SUMMARISE SUPPORTING DOCUMENTS (each its own call) ─────
        # Deliberately NOT combined into the main extraction call above —
        # see the comment there. Each supporting doc gets its own small,
        # focused summarisation call (short max_tokens, asks for a brief
        # note rather than full extraction), with a pause between calls so
        # token usage is spread across the per-minute window rather than
        # concentrated in one request.
        supporting_summaries = []
        if supporting_docs_b64:
            for i, doc in enumerate(supporting_docs_b64):
                progress(f'Reading supporting document {i+1} of {len(supporting_docs_b64)}: {doc["name"]}...', 6)
                summary_prompt = (
                    f'This is a supporting context document called "{doc["name"]}", attached alongside '
                    'a primary client brief for a hospitality interior design proposal. Read it and write '
                    'a concise note (4-8 sentences) covering only information relevant to a design proposal: '
                    'spaces, tiers, capacities, budgets, constraints, requirements, names, dates. Skip anything '
                    'not relevant to scoping interior design work. Do not invent information not present in the '
                    'document. Return plain text only, no markdown, no headers.'
                )
                doc_msg_content = [
                    {'type': 'document', 'source': {'type': 'base64', 'media_type': 'application/pdf', 'data': doc['b64']}},
                    {'type': 'text', 'text': summary_prompt},
                ]
                doc_resp = None
                for attempt in range(3):
                    try:
                        if attempt > 0:
                            wait = [0, 25, 45][attempt]
                            progress(f'Rate limit — retrying {doc["name"]} in {wait}s...', 6)
                            time.sleep(wait)
                        doc_resp = client.messages.create(
                            model='claude-sonnet-4-6',
                            max_tokens=400,
                            messages=[{'role': 'user', 'content': doc_msg_content}]
                        )
                        break
                    except anthropic.RateLimitError:
                        if attempt == 2:
                            doc_resp = None
                    except Exception:
                        doc_resp = None
                        break
                if doc_resp is not None:
                    summary_text = doc_resp.content[0].text.strip()
                    supporting_summaries.append(f'[{doc["name"]}]: {summary_text}')
                else:
                    supporting_summaries.append(
                        f'[{doc["name"]}]: Could not be read — rate limit or processing error. '
                        'Review this document manually before sending the proposal.'
                    )
                # Brief pause between supporting-doc calls, and before the
                # next stage of the pipeline, to spread token usage over time
                # rather than bursting several calls back to back.
                if i < len(supporting_docs_b64) - 1:
                    time.sleep(8)

        if supporting_summaries:
            meta['supporting_context'] = '\n\n'.join(supporting_summaries)
        update_job(job_id, meta=meta)
        raw_spaces = ex.get('spaces', [])
        spaces_text = '\n'.join(
            f"- {s.get('name','?')}"
            + (f" | {s.get('tier','')}" if s.get('tier') else '')
            + (f" | Level {s.get('level','')}" if s.get('level') else '')
            + (f" | Capacity {s.get('capacity','')}" if s.get('capacity') else '')
            + (f" | Budget {s.get('budget','')}" if s.get('budget') else '')
            + (f" — {s.get('notes','')}" if s.get('notes') else '')
            for s in raw_spaces
        ) or meta.get('scope', 'Not listed')

        # Inject prior work context if provided
        if prior_work:
            meta['continuation'] = 'yes'
            meta['prior_stages_completed'] = (meta.get('prior_stages_completed','') + ' ' + prior_work).strip()

        progress(f'Brief read — {meta["client"] or "client"} / {meta["venue"] or "project"}', 10)

        # ── PAUSE FOR REVIEW ─────────────────────────────────────────────────
        # Store everything needed for generation, then pause so the user can
        # confirm the four key fields before writing starts. The /confirm
        # endpoint resumes from here when the user clicks Generate.
        update_job(job_id,
                   status='awaiting_review',
                   meta=meta,
                   spaces_text=spaces_text)
        return  # Thread ends here — /confirm will start a new thread

        # ── STEP 2: RESEARCH ─────────────────────────────────────────────────
        progress('Researching the client...', 15)
        time.sleep(12)  # Let rate limit recover after extraction

        contact_str = meta.get("contact","") or ""
        org_str = meta.get("client","") or meta.get("venue","") or ""
        if not contact_str and not org_str:
            update_job(job_id, intel={})
        else:
            research_prompt = (
                f'Research {contact_str}{(" at " + org_str) if org_str else ""} for a hospitality design agency pitch. '
                f'Find publicly available information about this person and organisation. '
                'Return ONLY valid JSON:\n'
                '{"contact_profile":"2-3 sentences about the person","org_context":"2-3 sentences on the organisation right now","'
                'why_now":"why this brief likely exists","ambitions":"their strategic goals","confidence":"high|medium|low"}'
            )
            try:
                resp2 = client.messages.create(
                    model='claude-sonnet-4-6',
                    max_tokens=800,
                    tools=[{'type': 'web_search_20250305', 'name': 'web_search'}],
                    messages=[{'role': 'user', 'content': research_prompt}]
                )
                txt2 = ' '.join(b.text for b in resp2.content if hasattr(b, 'text'))
                m2 = re.search(r'\{[\s\S]*\}', txt2)
                raw_intel = json.loads(m2.group(0)) if m2 else {}
                clean_intel = {k: strip_html(v) if isinstance(v, str) else v for k, v in raw_intel.items()}
                update_job(job_id, intel=clean_intel)
            except Exception:
                update_job(job_id, intel={})

        progress('Client research complete', 20)

        # ── STEP 3: GENERATE SECTIONS ────────────────────────────────────────
        ctx = build_context(meta, spaces_text)
        sections = []
        total = len(SECTIONS)
        # Wider gap when supporting docs were processed earlier in this same
        # job — those calls already consumed part of the per-minute token
        # budget before this loop even starts, so the default 7s gap isn't
        # enough slack on lower usage tiers. This doesn't remove the
        # underlying ceiling (see console.anthropic.com/settings/limits) but
        # reduces how often a complex, multi-document brief trips it.
        GAP = 14 if supporting_docs_b64 else 7

        for i, (sid, label, prompt_tpl) in enumerate(SECTIONS):
            pct = 20 + int((i / total) * 65)
            progress(f'Writing: {label} ({i+1} of {total})', pct)

            prompt = prompt_tpl.format(
                contact=meta.get('contact', 'the contact'),
                client=meta.get('client', 'the client'),
                ctx=ctx
            )

            # Multi-space briefs (e.g. 8+ named lounges/boxes at a stadium)
            # push every stage section — not just stage456 — toward covering
            # many sub-stages in one response. An 800-token budget tuned for a
            # single-space brief truncates Deliverables (or whichever section
            # the model writes last) once enough spaces are in scope. Scale
            # up by space count for any stage section, with stage456 getting
            # an extra margin since it covers three RIBA stages at once.
            num_spaces = len(raw_spaces) if raw_spaces else 0
            is_stage_section = sid in ('stage1', 'stage2', 'stage3', 'stage456')
            if is_stage_section and num_spaces >= 5:
                section_max_tokens = 2400 if sid == 'stage456' else 1700
            elif sid == 'stage456':
                section_max_tokens = 1900
            else:
                section_max_tokens = 800

            for attempt in range(4):
                try:
                    if attempt > 0:
                        wait = [0, 35, 55, 80][attempt]
                        progress(f'Rate limit — retrying {label} in {wait}s...', pct)
                        time.sleep(wait)

                    resp3 = client.messages.create(
                        model='claude-sonnet-4-6',
                        max_tokens=section_max_tokens,
                        system=SYSTEM_PROMPT,
                        messages=[{'role': 'user', 'content': prompt}]
                    )
                    sec = {'id': sid, 'heading': label, 'body': resp3.content[0].text.strip()}
                    sections.append(sec)
                    append_section(job_id, sec)
                    break
                except anthropic.RateLimitError:
                    if attempt == 3:
                        sec = {'id': sid, 'heading': label, 'body': '[Could not generate — add manually]'}
                        sections.append(sec)
                        append_section(job_id, sec)
                except Exception as e:
                    sec = {'id': sid, 'heading': label, 'body': f'[Error: {str(e)[:80]}]'}
                    sections.append(sec)
                    append_section(job_id, sec)
                    break

            if i < total - 1:
                time.sleep(GAP)

        update_job(job_id, sections=sections)
        progress('All sections written', 85)

        # ── STEP 4: BUILD PPTX ───────────────────────────────────────────────
        progress('Building PowerPoint...', 88)
        try:
            import tempfile
            pptx_dir = tempfile.mkdtemp(prefix='2020_out_')
            pptx_path = os.path.join(pptx_dir, 'proposal.pptx')
            build_pptx_clean(sections, meta, pptx_path)
            if not os.path.exists(pptx_path):
                raise FileNotFoundError('PowerPoint was not created')
            update_job(job_id, pptx_path=pptx_path, status='done')
            progress('Done — click Download PowerPoint or Word Doc', 100)
            log_usage_event(
                job_id, 'completed',
                client=meta.get('client', ''), venue=meta.get('venue', ''),
                brief_type=meta.get('brief_type', ''), is_riba=meta.get('is_riba', ''),
                pptx_ready=True,
                duration_seconds=round(time.time() - pipeline_start_ts, 1),
                num_supporting_docs=len(supporting_docs_b64),
            )
        except Exception as pptx_err:
            import traceback
            err_detail = traceback.format_exc()
            update_job(job_id, status='done', pptx_error=str(pptx_err), pptx_traceback=err_detail[-500:])
            progress(f'Sections complete. PowerPoint failed: {pptx_err}', 100)
            log_usage_event(
                job_id, 'pptx_failed',
                client=meta.get('client', ''), venue=meta.get('venue', ''),
                brief_type=meta.get('brief_type', ''), is_riba=meta.get('is_riba', ''),
                pptx_ready=False, error=str(pptx_err)[:200],
                duration_seconds=round(time.time() - pipeline_start_ts, 1),
                num_supporting_docs=len(supporting_docs_b64),
            )

    except Exception as e:
        import traceback
        # Even on pipeline error, mark done if we have sections
        job = load_job(job_id) or {}
        meta_for_log = job.get('meta', {})
        if job.get('sections'):
            update_job(job_id, status='done', error=str(e))
            log_usage_event(
                job_id, 'completed_with_error',
                client=meta_for_log.get('client', ''), venue=meta_for_log.get('venue', ''),
                error=str(e)[:200],
                duration_seconds=round(time.time() - pipeline_start_ts, 1),
                num_supporting_docs=len(supporting_docs_b64),
            )
        else:
            update_job(job_id, status='error', error=str(e))
            log_usage_event(
                job_id, 'failed',
                client=meta_for_log.get('client', ''), venue=meta_for_log.get('venue', ''),
                error=str(e)[:200],
                duration_seconds=round(time.time() - pipeline_start_ts, 1),
                num_supporting_docs=len(supporting_docs_b64),
            )
        progress(f'Error: {e}', None)


def run_generation(job_id):
    """Runs research → sections → PPTX for a job that has already completed
    extraction and been confirmed via /confirm. Reads meta and spaces_text
    from the saved job rather than re-running extraction."""
    job = load_job(job_id) or {}
    meta = job.get('meta', {})
    spaces_text = job.get('spaces_text', meta.get('scope', 'Not listed'))
    pipeline_start_ts = job.get('pipeline_start_ts', time.time())
    supporting_docs_count = job.get('supporting_docs_b64_count', 0)
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)

    def progress(msg, pct=None):
        append_progress(job_id, msg, pct)

    update_job(job_id, status='running')
    try:
        # ── RESEARCH ─────────────────────────────────────────────────────────
        progress('Researching the client...', 15)
        time.sleep(12)  # Let rate limit recover after extraction

        contact_str = meta.get('contact', '') or ''
        org_str = meta.get('client', '') or meta.get('venue', '') or ''
        if not contact_str and not org_str:
            update_job(job_id, intel={})
        else:
            sector_str = meta.get('sector', '')
            research_prompt = (
                f'Research {contact_str}{(" at " + org_str) if org_str else ""} for a design agency pitch. '
                f'Organisation sector: {sector_str}. '
                f'Find publicly available information about this person and organisation. '
                'Return ONLY valid JSON:\n'
                '{"contact_profile":"2-3 sentences about the person","org_context":"2-3 sentences on the organisation right now",'
                '"why_now":"why this brief likely exists","ambitions":"their strategic goals","confidence":"high|medium|low"}'
            )
            try:
                resp2 = client.messages.create(
                    model='claude-sonnet-4-6',
                    max_tokens=800,
                    tools=[{'type': 'web_search_20250305', 'name': 'web_search'}],
                    messages=[{'role': 'user', 'content': research_prompt}]
                )
                txt2 = ' '.join(b.text for b in resp2.content if hasattr(b, 'text'))
                m2 = re.search(r'\{[\s\S]*\}', txt2)
                raw_intel = json.loads(m2.group(0)) if m2 else {}
                clean_intel = {k: strip_html(v) if isinstance(v, str) else v for k, v in raw_intel.items()}
                update_job(job_id, intel=clean_intel)
            except Exception:
                update_job(job_id, intel={})

        progress('Client research complete', 20)
        ctx = build_context(meta, spaces_text)

        # ── SECTIONS ─────────────────────────────────────────────────────────
        # Select the right template based on confirmed project type
        pt = meta.get('project_type', 'hospitality')
        active_sections = SECTIONS_GRAPHICS if pt in ('graphics_brand', 'other') else SECTIONS
        num_spaces = len(job.get('spaces_text', '').split('\n')) if job.get('spaces_text') else 0
        GAP = 14 if supporting_docs_count > 0 else 7
        sections = []
        total = len(active_sections)

        for i, (sid, label, prompt_tpl) in enumerate(active_sections):
            pct = 20 + int((i / total) * 60)
            progress(f'Writing {label}...', pct)

            is_stage = sid in ('stage1', 'stage2', 'stage3', 'stage456')
            if is_stage and num_spaces >= 5:
                section_max_tokens = 2400 if sid == 'stage456' else 1700
            elif sid == 'stage456':
                section_max_tokens = 1900
            else:
                section_max_tokens = 800

            prompt = prompt_tpl.format(
                contact=meta.get('contact', 'the contact'),
                client=meta.get('client', 'the client'),
                ctx=ctx
            )

            for attempt in range(4):
                try:
                    if attempt > 0:
                        wait = [0, 35, 55, 80][attempt]
                        progress(f'Rate limit — retrying {label} in {wait}s...', pct)
                        time.sleep(wait)
                    resp3 = client.messages.create(
                        model='claude-sonnet-4-6',
                        max_tokens=section_max_tokens,
                        system=SYSTEM_PROMPT,
                        messages=[{'role': 'user', 'content': prompt}]
                    )
                    sec = {'id': sid, 'heading': label, 'body': resp3.content[0].text.strip()}
                    sections.append(sec)
                    append_section(job_id, sec)
                    break
                except anthropic.RateLimitError:
                    if attempt == 3:
                        sec = {'id': sid, 'heading': label, 'body': '[Could not generate — add manually]'}
                        sections.append(sec)
                        append_section(job_id, sec)
                    print(f'SECTION ERROR [{sid}]: {str(e)}', flush=True)
                except Exception as e:
                    sec = {'id': sid, 'heading': label, 'body': f'[Error: {str(e)}]'}
                    sections.append(sec)
                    append_section(job_id, sec)
                    break

            if i < total - 1:
                time.sleep(GAP)

        update_job(job_id, sections=sections)
        progress('All sections written', 85)

        # ── BUILD PPTX ───────────────────────────────────────────────────────
        progress('Building PowerPoint...', 88)
        try:
            import tempfile as tf
            pptx_dir = tf.mkdtemp(prefix='2020_out_')
            pptx_path = os.path.join(pptx_dir, 'proposal.pptx')
            build_pptx_clean(sections, meta, pptx_path)
            if not os.path.exists(pptx_path):
                raise FileNotFoundError('PowerPoint was not created')
            update_job(job_id, pptx_path=pptx_path, status='done')
            progress('Done — click Download PowerPoint or Word Doc', 100)
            log_usage_event(job_id, 'completed',
                client=meta.get('client',''), venue=meta.get('venue',''),
                brief_type=meta.get('brief_type',''), project_type=meta.get('project_type',''),
                pptx_ready=True,
                duration_seconds=round(time.time() - pipeline_start_ts, 1))
        except Exception as pptx_err:
            update_job(job_id, status='done', pptx_error=str(pptx_err))
            progress(f'Sections complete. PowerPoint failed: {pptx_err}', 100)

    except Exception as e:
        job2 = load_job(job_id) or {}
        if job2.get('sections'):
            update_job(job_id, status='done', error=str(e))
        else:
            update_job(job_id, status='error', error=str(e))
        progress(f'Error: {e}', None)


# ── ROUTES ────────────────────────────────────────────────────────────────────
INDEX_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>20.20 Proposal Generator</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
:root{
  --nv:#1B2340;--gd:#C9A84C;--rd:#E97132;
  --bg:#F5F4F1;--white:#fff;--bd:#E0DED8;
  --tx:#1A1A1A;--tx2:#666;--r:8px;--rl:14px
}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Arial,sans-serif;
  background:var(--bg);color:var(--tx);min-height:100vh}

/* NAV */
nav{background:var(--nv);padding:0 2rem;display:flex;align-items:center;
  justify-content:space-between;height:54px;position:sticky;top:0;z-index:100}
.logo{display:flex;align-items:center;gap:10px}
.logo-mark{font-size:13px;font-weight:700;color:var(--white);line-height:1.1;letter-spacing:-0.5px}
.logo-mark span{color:var(--gd)}
.logo-name{font-size:13px;color:rgba(255,255,255,.55);border-left:1px solid rgba(255,255,255,.2);padding-left:10px}
.nav-status{font-size:12px;color:rgba(255,255,255,.45)}

/* LAYOUT */
.page{max-width:860px;margin:0 auto;padding:2rem 1.5rem 4rem}

/* PANELS */
.panel{background:var(--white);border-radius:var(--rl);margin-bottom:1.25rem;overflow:hidden;
  box-shadow:0 1px 4px rgba(0,0,0,.06)}
.panel-head{padding:1rem 1.25rem;border-bottom:1px solid var(--bd);display:flex;
  align-items:center;justify-content:space-between}
.panel-head h2{font-size:15px;font-weight:600;color:var(--nv)}
.panel-head .step-badge{font-size:10px;font-weight:700;text-transform:uppercase;
  letter-spacing:.08em;color:var(--tx2);background:var(--bg);
  padding:3px 10px;border-radius:20px}
.panel-body{padding:1.25rem}

/* FORM */
.tab-row{display:flex;border:1px solid var(--bd);border-radius:var(--r);overflow:hidden;margin-bottom:1rem}
.tab-btn{flex:1;padding:8px;font-size:13px;font-weight:600;border:none;cursor:pointer;
  font-family:inherit;transition:all .15s}
.tab-btn.active{background:var(--nv);color:var(--white)}
.tab-btn.inactive{background:var(--bg);color:var(--tx2);border-left:1px solid var(--bd)}
.field-label{display:block;font-size:12px;font-weight:600;margin-bottom:5px;color:var(--tx2)}
input[type=file]{display:block;width:100%;font-size:13px;padding:8px;
  border:1px solid var(--bd);border-radius:var(--r);background:var(--bg);cursor:pointer;font-family:inherit}
textarea{width:100%;font-size:13px;padding:10px;border:1px solid var(--bd);
  border-radius:var(--r);font-family:inherit;resize:vertical;min-height:160px;
  line-height:1.5;outline:none;transition:border-color .15s}
textarea:focus{border-color:var(--nv)}

/* BUTTONS */
.btn{display:inline-flex;align-items:center;gap:6px;padding:9px 20px;border:none;
  border-radius:var(--r);font-size:13px;font-weight:600;cursor:pointer;
  font-family:inherit;transition:opacity .15s}
.btn:hover{opacity:.88}
.btn-primary{background:var(--nv);color:var(--white)}
.btn-gold{background:var(--gd);color:var(--nv)}
.btn-outline{background:transparent;border:1px solid var(--bd);color:var(--tx)}
.btn:disabled{opacity:.4;cursor:not-allowed}

/* PROGRESS */
.progress-wrap{margin:1rem 0}
.progress-bar-bg{height:5px;background:var(--bd);border-radius:3px;overflow:hidden;margin-bottom:.5rem}
.progress-bar-fill{height:100%;background:var(--nv);border-radius:3px;transition:width .5s}
.progress-msg{font-size:12px;color:var(--tx2);text-align:center;min-height:1.2em}

/* SECTIONS */
.section-card{border:1px solid var(--bd);border-radius:var(--r);overflow:hidden;margin-bottom:10px}
.section-head{display:flex;align-items:center;justify-content:space-between;
  padding:.6rem 1rem;background:var(--nv);cursor:pointer;user-select:none}
.section-head-title{font-size:13px;font-weight:600;color:var(--white)}
.section-copy-btn{font-size:11px;color:rgba(255,255,255,.55);padding:2px 8px;
  background:rgba(255,255,255,.1);border:none;border-radius:4px;cursor:pointer;font-family:inherit}
.section-copy-btn:hover{background:rgba(255,255,255,.2);color:var(--white)}
.section-body{padding:.75rem 1rem;background:var(--white)}
.section-body textarea{min-height:80px;border:none;padding:0;background:transparent;
  font-size:13px;line-height:1.7;resize:vertical;outline:none;color:var(--tx)}

/* META */
.meta-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-bottom:1rem}
.meta-card{background:var(--bg);border:1px solid var(--bd);border-radius:var(--r);padding:.75rem 1rem}
.meta-label{font-size:10px;font-weight:700;text-transform:uppercase;letter-spacing:.06em;
  color:var(--tx2);margin-bottom:3px}
.meta-value{font-size:13px;color:var(--tx)}
.intel-card{background:var(--bg);border:1px solid var(--bd);border-radius:var(--r);padding:1rem}
.intel-row{padding:.5rem 0;border-bottom:1px solid var(--bd);font-size:13px;line-height:1.5}
.intel-row:last-child{border:none}
.intel-lbl{font-size:10px;font-weight:700;text-transform:uppercase;letter-spacing:.06em;
  color:var(--gd);margin-bottom:2px}

/* ACTIONS BAR */
.actions-bar{display:flex;gap:8px;flex-wrap:wrap;padding-top:1rem;
  border-top:1px solid var(--bd);margin-top:1rem}

/* HIDDEN */
.hidden{display:none}

/* PILL */
.pill{display:inline-block;font-size:10px;font-weight:700;text-transform:uppercase;
  letter-spacing:.06em;padding:2px 10px;border-radius:20px}
.pill-running{background:#EAF3DE;color:#3B6D11}
.pill-done{background:#E6F1FB;color:#185FA5}
.pill-error{background:#FCEBEB;color:#A32D2D}

/* ERROR */
.error-box{background:#FCEBEB;border:1px solid #F09595;border-radius:var(--r);
  padding:.75rem 1rem;font-size:13px;color:#A32D2D;margin-top:.75rem}

@media(max-width:600px){
  .meta-grid{grid-template-columns:1fr}
  .page{padding:1rem .75rem 3rem}
}
</style>
</head>
<body>

<nav>
  <div class="logo">
    <div class="logo-mark">20<br><span>20</span></div>
    <div class="logo-name">Proposal Generator</div>
  </div>
  <div class="nav-status" id="nav-status"></div>
</nav>

<div class="page">

  <!-- STEP 1: INPUT -->
  <div class="panel" id="panel-input">
    <div class="panel-head">
      <h2>Add the brief</h2>
      <span class="step-badge">Step 1</span>
    </div>
    <div class="panel-body">
      <p style="font-size:13px;color:var(--tx2);margin-bottom:1rem;line-height:1.5">
        Upload a PDF brief or paste the text. The tool reads it, researches the client,
        writes the proposal sections, and builds a branded PowerPoint — ready to review and download.
      </p>

      <div class="tab-row">
        <button class="tab-btn active" id="tab-pdf" onclick="switchTab('pdf')">↑ Upload PDF</button>
        <button class="tab-btn inactive" id="tab-text" onclick="switchTab('text')">Paste text</button>
      </div>

      <div id="panel-pdf">
        <label class="field-label">Select PDF brief</label>
        <input type="file" id="brief-pdf" accept=".pdf">
        <p style="font-size:11px;color:var(--tx2);margin-top:.4rem">
          Presentations, ITTs, emails saved as PDF — anything works
        </p>
      </div>

      <div id="panel-text" class="hidden">
        <label class="field-label">Paste brief text</label>
        <textarea id="brief-text" placeholder="Paste the brief here — email, copied PDF text, ITT, meeting notes..."></textarea>
      </div>

      <div style="margin-top:.75rem;padding:.75rem;background:var(--bg);border:1px solid var(--bd);border-radius:var(--r)">
        <div style="display:flex;align-items:center;gap:.5rem;margin-bottom:.4rem">
          <label class="field-label" style="margin:0">Continuation of prior work?</label>
          <label style="display:flex;align-items:center;gap:4px;font-size:12px;color:var(--tx2);cursor:pointer">
            <input type="checkbox" id="prior-work-toggle" onchange="togglePriorWork()">
            Yes — there is prior context to include
          </label>
        </div>
        <div id="prior-work-panel" style="display:none;margin-top:.5rem">
          <label class="field-label">Briefly describe the prior work (stages completed, key decisions made)</label>
          <textarea id="prior-work-text" rows="2" placeholder="e.g. Stages 1-3 completed in 2024. Concept was approved in March. Client likes the maroon palette but wants to revisit the bar layout in the Gold lounge..." style="width:100%;padding:8px;border:1px solid var(--bd);border-radius:var(--r);font-family:inherit;font-size:13px;resize:vertical"></textarea>
        </div>
      </div>

      <div style="margin-top:.75rem;padding:.75rem;background:var(--bg);border:1px solid var(--bd);border-radius:var(--r)">
        <div style="display:flex;align-items:center;gap:.5rem;margin-bottom:.4rem">
          <label class="field-label" style="margin:0">Supporting documents?</label>
          <label style="display:flex;align-items:center;gap:4px;font-size:12px;color:var(--tx2);cursor:pointer">
            <input type="checkbox" id="supporting-docs-toggle" onchange="toggleSupportingDocs()">
            Yes — add extra context documents
          </label>
        </div>
        <div id="supporting-docs-panel" style="display:none;margin-top:.5rem">
          <label class="field-label">Upload supporting PDFs (briefing decks, RFP appendices, prior reports, brand guidelines, etc.)</label>
          <input type="file" id="supporting-docs-input" accept=".pdf" multiple>
          <p style="font-size:11px;color:var(--tx2);margin-top:.4rem">
            Up to 3 files, PDF only. These are read alongside the main brief in a single pass, so keep them focused — every extra page adds to one request and can hit a rate limit on complex briefs. The main brief still takes priority if anything conflicts.
          </p>
          <div id="supporting-docs-list" style="margin-top:.4rem;font-size:12px;color:var(--tx2)"></div>
        </div>
      </div>

      <div id="submit-error" class="error-box hidden"></div>

      <div style="margin-top:1rem">
        <button class="btn btn-primary" id="submit-btn" onclick="submitBrief()">
          Generate proposal →
        </button>
      </div>
    </div>
  </div>

  <!-- STEP 2: PROGRESS -->
  <div class="panel hidden" id="panel-progress">
    <div class="panel-head">
      <h2>Generating</h2>
      <span class="pill pill-running" id="status-pill">Running</span>
    </div>
    <div class="panel-body">
      <div class="progress-wrap">
        <div class="progress-bar-bg"><div class="progress-bar-fill" id="prog-bar" style="width:0%"></div></div>
        <div class="progress-msg" id="prog-msg">Starting...</div>
      </div>
      <p style="font-size:12px;color:var(--tx2);line-height:1.5">
        Writing each section with a short pause between calls to stay within API rate limits.
        This takes around 90 seconds. Sections appear below as they complete.
      </p>
    </div>
  </div>

  <!-- REVIEW PANEL — appears after extraction, before generation -->
  <div class="panel hidden" id="panel-review">
    <div class="panel-head">
      <h2>Confirm before generating</h2>
      <span class="step-badge">Check and correct if needed</span>
    </div>
    <div class="panel-body">
      <div id="review-summary" style="padding:10px 14px;border-radius:6px;margin-bottom:1rem;font-size:13px;font-weight:500;background:#EAF3DE;border:1px solid #B7D89A;color:#3B6D11"></div>

      <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:1.25rem">

        <div>
          <label class="field-label">Project type</label>
          <select id="rv-project-type" class="t-sel">
            <option value="hospitality">Hospitality / Interior design</option>
            <option value="graphics_brand">Graphics &amp; Brand</option>
            <option value="other">Other</option>
          </select>
        </div>

        <div>
          <label class="field-label">Client</label>
          <input type="text" id="rv-client" class="t-sel" placeholder="Client name">
        </div>

        <div>
          <label class="field-label">Sector</label>
          <input type="text" id="rv-sector" class="t-sel" placeholder="e.g. Professional football club, Coffee brand">
        </div>

        <div>
          <label class="field-label">Scope of work</label>
          <input type="text" id="rv-scope" class="t-sel" placeholder="Plain-English description of what's in scope" style="text-overflow:ellipsis">
        </div>

      </div>

      <div id="rv-other-row" style="display:none;margin-bottom:1rem">
        <label class="field-label">Describe the project type</label>
        <input type="text" id="rv-other-desc" class="t-sel" placeholder="e.g. Production artwork for gift cards, Wayfinding for cruise vessel">
      </div>

      <button class="btn btn-primary" id="review-btn" onclick="submitReview()">
        Confirm and generate →
      </button>
      <span style="font-size:12px;color:var(--tx2);margin-left:1rem">Generation starts immediately.</span>
    </div>
  </div>

  <!-- TRIAGE NOTES -->
  <div class="panel hidden" id="panel-triage">
    <div class="panel-head"><h2>Triage</h2><span class="step-badge">Complete while generating</span></div>
    <div class="panel-body">
      <p style="font-size:12px;color:var(--tx2);margin-bottom:1rem">For internal use only — informs win likelihood score. Not sent to the AI.</p>

      <div style="display:grid;grid-template-columns:1fr 1fr 1fr 1fr;gap:10px;margin-bottom:.75rem">
        <div><label class="field-label">Pursue this brief?</label>
          <select id="t-pursue" class="t-sel" onchange="calcWin()">
            <option value="">— select —</option>
            <option value="3">Yes — full proposal</option>
            <option value="2">Yes — credentials only</option>
            <option value="1">Conditional</option>
            <option value="0">No</option>
          </select></div>
        <div><label class="field-label">Client status</label>
          <select id="t-client-status" class="t-sel" onchange="calcWin()">
            <option value="">— select —</option>
            <option value="3">Current client — ongoing</option>
            <option value="3">Returning client</option>
            <option value="2">Warm — previous contact</option>
            <option value="1">New — inbound approach</option>
            <option value="0">New — cold outreach</option>
          </select></div>
        <div><label class="field-label">Competitive pitch?</label>
          <select id="t-competitive" class="t-sel" onchange="calcWin()">
            <option value="">— select —</option>
            <option value="3">Direct appointment</option>
            <option value="2">2-3 agencies</option>
            <option value="1">4-5 agencies</option>
            <option value="0">Open tender (6+)</option>
          </select></div>
        <div><label class="field-label">Brief quality</label>
          <select id="t-brief-quality" class="t-sel" onchange="calcWin()">
            <option value="">— select —</option>
            <option value="3">Detailed — clear scope and budget</option>
            <option value="2">Good — scope clear, budget TBC</option>
            <option value="1">Outline — needs development</option>
            <option value="0">Vague — significant unknowns</option>
          </select></div>
      </div>

      <div style="display:grid;grid-template-columns:1fr 1fr 1fr 1fr;gap:10px;margin-bottom:.75rem">
        <div><label class="field-label">Resource available?</label>
          <select id="t-resource" class="t-sel" onchange="calcWin()">
            <option value="">— select —</option>
            <option value="2">Yes — team ready</option>
            <option value="1">Tight but manageable</option>
            <option value="1">Would need to juggle</option>
            <option value="0">No capacity currently</option>
          </select></div>
        <div><label class="field-label">Timescale realistic?</label>
          <select id="t-timescale" class="t-sel" onchange="calcWin()">
            <option value="">— select —</option>
            <option value="2">Yes — comfortable</option>
            <option value="1">Tight but achievable</option>
            <option value="0">Unrealistic as stated</option>
          </select></div>
        <div><label class="field-label">Design lead</label>
          <input type="text" id="t-lead" placeholder="Name" class="t-inp"></div>
        <div><label class="field-label">Creative lead</label>
          <input type="text" id="t-creative" placeholder="Name" class="t-inp"></div>
      </div>

      <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-bottom:.75rem">
        <div><label class="field-label">Other agencies shortlisted (if known)</label>
          <input type="text" id="t-competitors" placeholder="e.g. Bergman Interiors, Loop Creative" class="t-inp"></div>
        <div><label class="field-label">Key concerns or conditions</label>
          <input type="text" id="t-concerns" placeholder="e.g. Budget too low, timeline needs clarifying" class="t-inp"></div>
      </div>

      <!-- Win likelihood score -->
      <div id="win-score-panel" style="display:none;background:var(--nv);border-radius:var(--r);padding:.75rem 1rem;display:flex;align-items:center;gap:1rem">
        <div style="font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.06em;color:rgba(255,255,255,.5)">Win likelihood</div>
        <div id="win-score-value" style="font-size:28px;font-weight:700;color:#fff;font-family:inherit">—</div>
        <div id="win-score-label" style="font-size:13px;color:rgba(255,255,255,.7)"></div>
        <div style="flex:1"></div>
        <div id="win-score-bar-wrap" style="width:180px;height:8px;background:rgba(255,255,255,.15);border-radius:4px;overflow:hidden">
          <div id="win-score-bar" style="height:100%;border-radius:4px;transition:width .4s"></div>
        </div>
      </div>

    </div>
  </div>

  <!-- STEP 3: SECTIONS (appear during generation) -->
  <div class="panel hidden" id="panel-sections">
    <div class="panel-head">
      <h2>Proposal sections</h2>
      <span class="step-badge">Review and edit</span>
    </div>
    <div class="panel-body">
      <p style="font-size:12px;color:var(--tx2);margin-bottom:1rem">
        Each section is editable. Make any changes before downloading the PowerPoint.
      </p>
      <div id="sections-list"></div>
    </div>
  </div>

  <!-- STEP 4: CLIENT INTEL -->
  <div class="panel hidden" id="panel-intel">
    <div class="panel-head">
      <h2>Client intelligence</h2>
      <span class="step-badge">Verify before pitch</span>
    </div>
    <div class="panel-body" id="intel-body"></div>
  </div>

  <!-- STEP 5: ACTIONS -->
  <div class="panel hidden" id="panel-actions">
    <div class="panel-head">
      <h2>Download</h2>
      <span class="step-badge">Step 2</span>
    </div>
    <div class="panel-body">
      <p style="font-size:13px;color:var(--tx2);margin-bottom:1rem;line-height:1.5">
        The PowerPoint uses the 20.20 branded template with Filson Pro fonts, correct layouts and your client's colour.
        Image placeholders include specific direction for the creative team.
        Fees show [FEE: TBC] — apply the rate card before sending.
      </p>
      <div class="actions-bar">
        <button class="btn btn-gold" id="download-btn" onclick="downloadPPTX()">
          ↓ Download PowerPoint
        </button>
        <button class="btn btn-outline" id="download-docx-btn" onclick="downloadDocx()" style="display:none">
          ↓ Download Word Doc
        </button>
        <button class="btn btn-outline" onclick="rebuildAndDownload()">
          ↓ Rebuild from edited sections
        </button>
        <button class="btn btn-outline" onclick="resetAll()" style="margin-left:auto">
          New brief
        </button>
      </div>
      <div id="rebuild-status" style="font-size:12px;color:var(--tx2);margin-top:.5rem;display:none"></div>
    </div>
  </div>

</div><!-- .page -->

<script>
let activeTab = 'pdf';
let currentJobId = null;
let pollInterval = null;
let lastProgressLen = 0;
let currentMeta = {};
let jobDone = false;

function switchTab(t) {
  activeTab = t;
  document.getElementById('panel-pdf').classList.toggle('hidden', t !== 'pdf');
  document.getElementById('panel-text').classList.toggle('hidden', t !== 'text');
  document.getElementById('tab-pdf').className = 'tab-btn ' + (t === 'pdf' ? 'active' : 'inactive');
  document.getElementById('tab-text').className = 'tab-btn ' + (t === 'text' ? 'active' : 'inactive');
}

// ── REVIEW PANEL ──────────────────────────────────────────────────────────
function populateReviewPanel(data) {
  var r = data.review || {};
  var meta = data.meta || {};

  // Summary strip
  var summary = r.brief_summary || meta.brief_summary || 'Brief read — please confirm the details below before generating.';
  document.getElementById('review-summary').textContent = '✓  ' + summary;

  // Populate fields
  var ptSel = document.getElementById('rv-project-type');
  var pt = r.project_type || '';
  var matchedPt = false;
  for (var i = 0; i < ptSel.options.length; i++) {
    if (ptSel.options[i].value === pt) { ptSel.selectedIndex = i; matchedPt = true; break; }
  }
  if (!matchedPt && pt) {
    // Unknown type — show as Other with description
    ptSel.value = 'other';
    document.getElementById('rv-other-row').style.display = 'block';
    document.getElementById('rv-other-desc').value = pt;
  }

  document.getElementById('rv-client').value = r.client || meta.client || '';
  document.getElementById('rv-sector').value = r.sector || meta.sector || '';
  document.getElementById('rv-scope').value  = r.scope_plain || meta.scope_plain || meta.scope || '';

  // Show/hide other description on project type change
  ptSel.onchange = function() {
    document.getElementById('rv-other-row').style.display =
      ptSel.value === 'other' ? 'block' : 'none';
  };
}

async function submitReview() {
  var btn = document.getElementById('review-btn');
  btn.disabled = true;
  btn.textContent = 'Starting generation…';

  var ptVal = document.getElementById('rv-project-type').value;
  if (ptVal === 'other') {
    var desc = document.getElementById('rv-other-desc').value.trim();
    ptVal = desc || 'other';
  }

  var payload = {
    project_type: ptVal,
    client:       document.getElementById('rv-client').value.trim(),
    sector:       document.getElementById('rv-sector').value.trim(),
    scope_plain:  document.getElementById('rv-scope').value.trim(),
  };

  try {
    var resp = await fetch('/confirm/' + currentJobId, {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(payload)
    });
    var data = await resp.json();
    if (data.error) throw new Error(data.error);
    document.getElementById('panel-review').classList.add('hidden');
    pollInterval = setInterval(pollStatus, 2000);
  } catch(e) {
    btn.disabled = false;
    btn.textContent = 'Confirm and generate →';
    alert('Could not start generation: ' + e.message);
  }
}

function togglePriorWork() {
  var chk = document.getElementById('prior-work-toggle');
  document.getElementById('prior-work-panel').style.display = chk.checked ? 'block' : 'none';
}

function toggleSupportingDocs() {
  var chk = document.getElementById('supporting-docs-toggle');
  document.getElementById('supporting-docs-panel').style.display = chk.checked ? 'block' : 'none';
}

document.addEventListener('DOMContentLoaded', function() {
  var input = document.getElementById('supporting-docs-input');
  if (input) {
    input.addEventListener('change', function() {
      var list = document.getElementById('supporting-docs-list');
      var files = Array.from(input.files || []);
      if (files.length > 3) {
        list.innerHTML = '<span style="color:#A32D2D">Please select up to 3 files — only the first 3 will be used.</span>';
        files = files.slice(0, 3);
      } else if (files.length) {
        list.innerHTML = files.map(function(f) { return '\u2713 ' + f.name; }).join('<br>');
      } else {
        list.innerHTML = '';
      }
    });
  }
});

async function submitBrief() {
  var errEl = document.getElementById('submit-error');
  errEl.classList.add('hidden');

  var fd = new FormData();
  if (activeTab === 'pdf') {
    var f = document.getElementById('brief-pdf').files[0];
    if (!f) { errEl.textContent = 'Please select a PDF file first.'; errEl.classList.remove('hidden'); return; }
    fd.append('brief_pdf', f);
  } else {
    var txt = document.getElementById('brief-text').value.trim();
    if (!txt) { errEl.textContent = 'Please paste the brief text.'; errEl.classList.remove('hidden'); return; }
    fd.append('brief_text', txt);
  }

  // Add prior work context if provided
  var priorToggle = document.getElementById('prior-work-toggle');
  if (priorToggle && priorToggle.checked) {
    var priorTxt = document.getElementById('prior-work-text').value.trim();
    if (priorTxt) fd.append('prior_work_context', priorTxt);
  }

  // Add supporting documents if provided — up to 5 PDFs, read alongside
  // the main brief during extraction (main brief still takes priority).
  var docsToggle = document.getElementById('supporting-docs-toggle');
  if (docsToggle && docsToggle.checked) {
    var docsInput = document.getElementById('supporting-docs-input');
    var docFiles = Array.from((docsInput && docsInput.files) || []).slice(0, 3);
    docFiles.forEach(function(file) { fd.append('supporting_docs', file); });
  }

  document.getElementById('submit-btn').disabled = true;
  document.getElementById('submit-btn').textContent = 'Submitting...';

  try {
    var resp = await fetch('/submit', { method: 'POST', body: fd });
    var data = await resp.json();
    if (data.error) throw new Error(data.error);
    currentJobId = data.job_id;
    // Put job_id in URL so user can bookmark/debug
    window.history.replaceState(null, '', '/?job=' + data.job_id);
    showProgress();
    pollInterval = setInterval(pollStatus, 2000);
  } catch(e) {
    errEl.textContent = e.message;
    errEl.classList.remove('hidden');
    document.getElementById('submit-btn').disabled = false;
    document.getElementById('submit-btn').textContent = 'Generate proposal →';
  }
}

function showProgress() {
  document.getElementById('panel-progress').classList.remove('hidden');
  document.getElementById('nav-status').textContent = 'Generating… (job: ' + currentJobId + ')';
}

async function pollStatus() {
  if (!currentJobId) return;
  try {
    var resp = await fetch('/status/' + currentJobId);
    var data = await resp.json();

    // Update progress
    if (data.progress && data.progress.length > lastProgressLen) {
      var latest = data.progress[data.progress.length - 1];
      document.getElementById('prog-msg').textContent = latest.msg;
      if (latest.pct != null) {
        document.getElementById('prog-bar').style.width = latest.pct + '%';
      }
      lastProgressLen = data.progress.length;
    }

    // Show sections as they arrive
    if (data.sections && data.sections.length > 0) {
      document.getElementById('panel-sections').classList.remove('hidden');
      renderSections(data.sections);
    }

    // Review panel — shown after extraction, before generation starts
    if (data.status === 'awaiting_review') {
      clearInterval(pollInterval);
      pollInterval = null;
      populateReviewPanel(data);
      document.getElementById('panel-review').classList.remove('hidden');
      return;
    }

    // Store meta and reveal triage panel
    if (data.meta && data.meta.client) {
      currentMeta = data.meta;
      document.getElementById('panel-triage').classList.remove('hidden');
    }

    // Done
    if (data.status === 'done' || data.status === 'error') {
      clearInterval(pollInterval);
      jobDone = true;
      document.getElementById('prog-bar').style.width = '100%';

      if (data.status === 'done') {
        document.getElementById('status-pill').textContent = 'Complete';
        document.getElementById('status-pill').className = 'pill pill-done';
        document.getElementById('nav-status').textContent = data.pptx_ready ? 'Ready to download' : 'Sections ready';
        document.getElementById('prog-msg').textContent = data.pptx_ready ? 'Complete' : (data.pptx_error || 'Sections complete');
      } else {
        document.getElementById('status-pill').textContent = 'Error';
        document.getElementById('status-pill').className = 'pill pill-error';
        document.getElementById('nav-status').textContent = 'Error';
        document.getElementById('prog-msg').textContent = data.error || 'Unknown error';
      }

      // Show intel if available
      if (data.intel && Object.keys(data.intel).length) {
        renderIntel(data.intel, data.meta);
      }
      // Always show actions panel if we have sections
      if (data.sections && data.sections.length > 0) {
        var actionsPanel = document.getElementById('panel-actions');
        actionsPanel.classList.remove('hidden');
        // Update download button based on pptx status
        var dlBtn = document.getElementById('download-btn');
        if (data.pptx_ready) {
          dlBtn.disabled = false;
          dlBtn.textContent = '↓ Download PowerPoint';
          document.getElementById('download-docx-btn').style.display = 'inline-block';
        } else {
          dlBtn.disabled = true;
          dlBtn.textContent = 'PowerPoint unavailable — use Rebuild';
          document.getElementById('download-docx-btn').style.display = 'inline-block';
          // Show error detail
          var rs = document.getElementById('rebuild-status');
          rs.style.display = 'block';
          rs.style.color = '#A32D2D';
          rs.textContent = data.pptx_error
            ? 'PowerPoint build failed: ' + data.pptx_error + '. Try "Rebuild from edited sections".'
            : 'PowerPoint not built. Try "Rebuild from edited sections".';
        }
        actionsPanel.scrollIntoView({ behavior: 'smooth', block: 'start' });
      }
    }

  } catch(e) {
    console.error('Poll error:', e);
  }
}

function renderSections(sections) {
  var list = document.getElementById('sections-list');
  sections.forEach(function(sec, i) {
    var existing = document.getElementById('sec-card-' + sec.id);
    if (existing) {
      // Update textarea if user hasn't edited it
      var ta = existing.querySelector('textarea');
      if (ta && ta.dataset.pristine !== 'false') {
        ta.value = sec.body;
        ta.style.height = 'auto';
        ta.style.height = ta.scrollHeight + 'px';
      }
      return;
    }
    var card = document.createElement('div');
    card.className = 'section-card';
    card.id = 'sec-card-' + sec.id;
    card.innerHTML =
      '<div class="section-head" onclick="toggleSection(this)">' +
        '<span class="section-head-title">' + sec.heading + '</span>' +
        '<button class="section-copy-btn" onclick="copySec(event,\\'' + sec.id + '\\')">Copy</button>' +
      '</div>' +
      '<div class="section-body">' +
        '<textarea id="sec-ta-' + sec.id + '" onchange="this.dataset.pristine=\\'false\\'" ' +
          'oninput="this.style.height=\\'auto\\';this.style.height=this.scrollHeight+\\'px\\'">' +
          escHtml(sec.body) + '</textarea>' +
      '</div>';
    list.appendChild(card);
    var ta = card.querySelector('textarea');
    ta.dataset.pristine = 'true';
    setTimeout(function(){ ta.style.height='auto'; ta.style.height=ta.scrollHeight+'px'; }, 50);
    card.scrollIntoView({ behavior: 'smooth', block: 'nearest' });
  });
}

function toggleSection(head) {
  var body = head.nextElementSibling;
  body.style.display = body.style.display === 'none' ? 'block' : 'none';
}

function copySec(e, sid) {
  e.stopPropagation();
  var ta = document.getElementById('sec-ta-' + sid);
  if (!ta) return;
  var btn = e.target;
  navigator.clipboard.writeText(ta.value).then(function() {
    btn.textContent = 'Copied ✓';
    setTimeout(function(){ btn.textContent = 'Copy'; }, 2000);
  });
}

function renderIntel(intel, meta) {
  document.getElementById('panel-intel').classList.remove('hidden');
  var rows = [];
  if (intel.contact_profile) rows.push(['Contact', intel.contact_profile]);
  if (intel.org_context)     rows.push(['Organisation right now', intel.org_context]);
  if (intel.why_now)         rows.push(['Why this brief exists', intel.why_now]);
  if (intel.ambitions)       rows.push(['Strategic ambitions', intel.ambitions]);

  document.getElementById('intel-body').innerHTML =
    '<div class="intel-card">' +
    rows.map(function(r) {
      return '<div class="intel-row"><div class="intel-lbl">' + r[0] + '</div>' + escHtml(r[1]) + '</div>';
    }).join('') +
    '<p style="font-size:11px;color:var(--tx2);margin-top:.75rem">Verify key facts before the pitch meeting.</p>' +
    '</div>';
}

// Triage input styles applied via class
(function() {
  var style = document.createElement('style');
  style.textContent = '.t-sel,.t-inp{width:100%;padding:7px;border:1px solid var(--bd);border-radius:var(--r);font-family:inherit;font-size:13px;background:var(--bg)} .t-sel:focus,.t-inp:focus{outline:none;border-color:var(--nv)}';
  document.head.appendChild(style);
})();

function calcWin() {
  var fields = ['t-pursue','t-client-status','t-competitive','t-brief-quality','t-resource','t-timescale'];
  var total = 0; var filled = 0; var max = 13; // 3+3+3+3+2+2 = 16 max but pursue=0 exits
  var pursue = document.getElementById('t-pursue');
  if (pursue && pursue.value === '0') {
    document.getElementById('win-score-value').textContent = 'Pass';
    document.getElementById('win-score-label').textContent = 'Decision: do not pursue';
    document.getElementById('win-score-bar').style.width = '0%';
    document.getElementById('win-score-bar').style.background = '#E53935';
    document.getElementById('win-score-panel').style.display = 'flex';
    return;
  }
  fields.forEach(function(id) {
    var el = document.getElementById(id);
    if (el && el.value !== '') { total += parseInt(el.value||0); filled++; }
  });
  if (filled < 2) { document.getElementById('win-score-panel').style.display = 'none'; return; }
  var pct = Math.round((total / 16) * 100);
  var label, colour;
  if (pct >= 75)      { label = 'Strong — prioritise this one'; colour = '#43A047'; }
  else if (pct >= 55) { label = 'Good — worth a full effort'; colour = '#C9A84C'; }
  else if (pct >= 35) { label = 'Marginal — credentials only?'; colour = '#E97132'; }
  else                { label = 'Low — consider declining'; colour = '#E53935'; }
  document.getElementById('win-score-value').textContent = pct + '%';
  document.getElementById('win-score-label').textContent = label;
  document.getElementById('win-score-bar').style.width = pct + '%';
  document.getElementById('win-score-bar').style.background = colour;
  document.getElementById('win-score-panel').style.display = 'flex';
}

function collectSections() {
  var secs = [];
  document.querySelectorAll('[id^="sec-ta-"]').forEach(function(ta) {
    var sid = ta.id.replace('sec-ta-', '');
    var card = document.getElementById('sec-card-' + sid);
    var heading = card ? card.querySelector('.section-head-title').textContent : sid;
    secs.push({ id: sid, heading: heading, body: ta.value });
  });
  return secs;
}

function downloadPPTX() {
  if (!currentJobId) return;
  window.location.href = '/download/' + currentJobId;
}

function downloadDocx() {
  if (!currentJobId) return;
  window.location.href = '/download-docx/' + currentJobId;
}

async function rebuildAndDownload() {
  var sections = collectSections();
  var st = document.getElementById('rebuild-status');
  st.style.display = 'block';
  st.textContent = 'Rebuilding PowerPoint from your edited sections...';

  try {
    var resp = await fetch('/rebuild', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ job_id: currentJobId, sections: sections, meta: currentMeta })
    });
    var data = await resp.json();
    if (data.error) throw new Error(data.error);
    st.textContent = 'Done — downloading...';
    setTimeout(function() { window.location.href = '/download/' + currentJobId; }, 500);
  } catch(e) {
    st.textContent = 'Error: ' + e.message;
    st.style.color = '#A32D2D';
  }
}

function resetAll() {
  clearInterval(pollInterval);
  currentJobId = null;
  lastProgressLen = 0;
  currentMeta = {};
  jobDone = false;
  document.getElementById('sections-list').innerHTML = '';
  document.getElementById('intel-body').innerHTML = '';
  document.getElementById('brief-pdf').value = '';
  document.getElementById('brief-text').value = '';
  document.getElementById('prog-bar').style.width = '0%';
  document.getElementById('prog-msg').textContent = 'Starting...';
  document.getElementById('submit-btn').disabled = false;
  document.getElementById('submit-btn').textContent = 'Generate proposal →';
  document.getElementById('nav-status').textContent = '';
  ['panel-progress','panel-review','panel-triage','panel-sections','panel-intel','panel-actions'].forEach(function(id) {
    document.getElementById(id).classList.add('hidden');
  });
  document.getElementById('submit-error').classList.add('hidden');
  document.getElementById('rebuild-status').style.display = 'none';
  window.scrollTo({ top: 0, behavior: 'smooth' });
}

function escHtml(s) {
  return (s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}
</script>
</body>
</html>
"""

@app.route('/health')
def health():
    """Diagnostic route — shows what the server can find."""
    import glob
    here = os.path.dirname(os.path.abspath(__file__))
    cwd = os.getcwd()
    files_here = os.listdir(here)
    files_cwd = os.listdir(cwd) if cwd != here else '(same as above)'
    return jsonify({
        'status': 'running',
        'api_key_set': bool(ANTHROPIC_KEY),
        'api_key_prefix': ANTHROPIC_KEY[:12] + '...' if ANTHROPIC_KEY else None,
        'template_path': TEMPLATE_PATH,
        'template_exists': os.path.exists(TEMPLATE_PATH),
        'here': here,
        'cwd': cwd,
        'files_in_app_dir': files_here,
        'files_in_cwd': files_cwd,
        'jobs_dir': JOBS_DIR,
        'jobs_dir_exists': os.path.exists(JOBS_DIR),
    })

@app.route('/usage')
def usage_dashboard():
    """Internal usage dashboard — shows every generation attempt logged via
    log_usage_event(). Protected by USAGE_DASHBOARD_KEY (set in Railway's
    environment variables); without it set, this route refuses all access.
    Visit /usage?key=YOUR_KEY to view. Visit /usage?key=YOUR_KEY&migrate=1
    once to backfill from any job files still on disk (best-effort — see
    migrate_existing_jobs_to_usage_log for why this can't recover jobs from
    before the last server restart)."""
    if not USAGE_DASHBOARD_KEY:
        return 'Usage dashboard is disabled — set USAGE_DASHBOARD_KEY in Railway environment variables to enable it.', 403
    if request.args.get('key') != USAGE_DASHBOARD_KEY:
        return 'Forbidden — missing or incorrect key.', 403

    migration_summary = None
    if request.args.get('migrate'):
        migration_summary = migrate_existing_jobs_to_usage_log()

    events = read_usage_log()
    events.sort(key=lambda e: e.get('ts', 0), reverse=True)

    total_started = sum(1 for e in events if e.get('event') == 'started')
    total_completed = sum(1 for e in events if e.get('event') in ('completed', 'completed_with_error'))
    total_failed = sum(1 for e in events if e.get('event') in ('failed', 'pptx_failed'))
    docs_used = sum(1 for e in events if e.get('event') == 'started' and e.get('num_supporting_docs', 0) > 0)

    def esc(s):
        return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    rows_html = []
    for e in events:
        event = e.get('event', '')
        colour = {'completed': '#3B6D11', 'started': '#666', 'failed': '#A32D2D',
                  'pptx_failed': '#A32D2D', 'completed_with_error': '#C9A84C'}.get(event, '#666')
        rows_html.append(f"""
            <tr>
              <td>{esc(e.get('date',''))}</td>
              <td><span style="color:{colour};font-weight:600">{esc(event)}</span></td>
              <td>{esc(e.get('client',''))}</td>
              <td>{esc(e.get('venue',''))}</td>
              <td>{esc(e.get('brief_type',''))}</td>
              <td>{esc(e.get('duration_seconds','')) }{'s' if e.get('duration_seconds') else ''}</td>
              <td>{esc(e.get('num_supporting_docs',''))}</td>
              <td style="max-width:280px;overflow:hidden;text-overflow:ellipsis;color:#A32D2D">{esc(e.get('error',''))}</td>
              <td style="color:#999;font-size:11px">{esc(e.get('job_id',''))}</td>
            </tr>""")

    migration_html = ''
    if migration_summary is not None:
        migration_html = (
            f'<div style="background:#EAF3DE;border:1px solid #B7D89A;border-radius:6px;'
            f'padding:10px 14px;margin-bottom:1rem;font-size:13px">'
            f'Migration ran: found {migration_summary["found"]} job file(s) on disk, '
            f'added {migration_summary["migrated"]} new event(s), '
            f'skipped {migration_summary["skipped"]} (already logged or unreadable). '
            f'Note: this can only recover jobs that survived up to the most recent server restart — '
            f'JOBS_DIR is ephemeral storage on Railway and is wiped on every redeploy.'
            f'</div>'
        )

    html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>Usage — 20.20 Proposal Generator</title>
<style>
body{{font-family:-apple-system,Arial,sans-serif;background:#F5F4F1;margin:0;padding:2rem;color:#1A1A1A}}
h1{{font-size:20px;margin-bottom:.25rem}}
.sub{{color:#666;font-size:13px;margin-bottom:1.5rem}}
.stats{{display:flex;gap:1rem;margin-bottom:1.5rem;flex-wrap:wrap}}
.stat{{background:#fff;border:1px solid #E0DED8;border-radius:8px;padding:.75rem 1.25rem;min-width:120px}}
.stat .n{{font-size:24px;font-weight:700;color:#1B2340}}
.stat .l{{font-size:11px;text-transform:uppercase;letter-spacing:.05em;color:#666}}
table{{width:100%;border-collapse:collapse;background:#fff;border-radius:8px;overflow:hidden;font-size:13px}}
th{{background:#1B2340;color:#fff;text-align:left;padding:8px 10px;font-size:11px;text-transform:uppercase;letter-spacing:.04em}}
td{{padding:7px 10px;border-bottom:1px solid #EEE}}
tr:hover{{background:#FAFAF8}}
a{{color:#1B2340}}
</style></head>
<body>
  <h1>Proposal generator — usage log</h1>
  <div class="sub">{len(events)} total events logged. <a href="?key={esc(USAGE_DASHBOARD_KEY)}&migrate=1">Run migration from existing job files</a></div>
  {migration_html}
  <div class="stats">
    <div class="stat"><div class="n">{total_started}</div><div class="l">Generations started</div></div>
    <div class="stat"><div class="n">{total_completed}</div><div class="l">Completed</div></div>
    <div class="stat"><div class="n">{total_failed}</div><div class="l">Failed</div></div>
    <div class="stat"><div class="n">{docs_used}</div><div class="l">Used supporting docs</div></div>
  </div>
  <table>
    <tr><th>Date</th><th>Event</th><th>Client</th><th>Venue</th><th>Brief type</th><th>Duration</th><th>Supp. docs</th><th>Error</th><th>Job ID</th></tr>
    {''.join(rows_html) if rows_html else '<tr><td colspan="9" style="text-align:center;color:#999;padding:2rem">No usage logged yet.</td></tr>'}
  </table>
</body></html>"""
    return html

def index():
    return INDEX_HTML

@app.route('/')
def index():
    return INDEX_HTML

@app.route('/submit', methods=['POST'])
def submit():
    if not ANTHROPIC_KEY:
        return jsonify({'error': 'API key not configured on server.'}), 500

    job_id = str(uuid.uuid4())[:8]

    try:
        save_job(job_id, {
            'status': 'running',
            'progress': [],
            'sections': [],
            'meta': {},
            'intel': {},
            'extracted': {},
            'pptx_path': None,
            'error': None,
        })

        pdf_b64 = None
        brief_text = None
        prior_work = request.form.get('prior_work_context', '')

        if 'brief_pdf' in request.files and request.files['brief_pdf'].filename:
            f = request.files['brief_pdf']
            pdf_b64 = base64.b64encode(f.read()).decode('ascii')
        elif request.form.get('brief_text'):
            brief_text = request.form.get('brief_text')
        else:
            return jsonify({'error': 'Please upload a PDF or paste the brief text.'}), 400

        # Supporting context documents — up to 3 extra PDFs read alongside the
        # main brief during extraction. The main brief still takes priority if
        # anything conflicts; these just add context the model wouldn't
        # otherwise see (RFP appendices, brand guidelines, prior reports, etc.)
        supporting_docs_b64 = []
        supporting_files = request.files.getlist('supporting_docs')[:3]
        for sf in supporting_files:
            if not sf or not sf.filename:
                continue
            if not sf.filename.lower().endswith('.pdf'):
                continue
            data = sf.read()
            if len(data) > 8 * 1024 * 1024:  # 8MB per supporting file
                continue
            supporting_docs_b64.append({
                'name': sf.filename,
                'b64': base64.b64encode(data).decode('ascii'),
            })

    except RequestEntityTooLarge:
        return jsonify({'error': (
            'The brief and supporting documents together are too large for one upload. '
            'Try fewer supporting documents, or smaller files, and submit again.'
        )}), 413
    except Exception as e:
        import traceback
        print('SUBMIT ERROR:', traceback.format_exc())
        return jsonify({'error': f'Could not process the upload: {str(e)[:200]}'}), 500

    t = threading.Thread(target=run_pipeline,
                          args=(job_id, pdf_b64, brief_text, prior_work, supporting_docs_b64),
                          daemon=True)
    t.start()

    return jsonify({'job_id': job_id})


@app.route('/confirm/<job_id>', methods=['POST'])
def confirm(job_id):
    """Accept the four confirmed review fields, merge into meta, start generation."""
    job = load_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404

    data = request.get_json() or {}
    meta = job.get('meta', {})

    # Merge confirmed values — only update fields that were sent
    for field in ('project_type', 'client', 'sector', 'scope_plain'):
        if field in data and data[field] is not None:
            meta[field] = data[field]
            if field == 'scope_plain':
                meta['scope'] = data[field]

    # Determine is_riba from confirmed project_type
    pt = meta.get('project_type', '')
    if pt and pt != 'hospitality':
        meta['is_riba'] = 'no'

    # Store pipeline start time if not already set
    if not job.get('pipeline_start_ts'):
        job['pipeline_start_ts'] = time.time()

    update_job(job_id, meta=meta, status='running',
               pipeline_start_ts=job.get('pipeline_start_ts', time.time()))

    t = threading.Thread(target=run_generation, args=(job_id,), daemon=True)
    t.start()
    return jsonify({'ok': True})

@app.route('/debug/<job_id>')
def debug(job_id):
    """Shows full job state for troubleshooting."""
    job = load_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found', 'jobs_dir': JOBS_DIR}), 404
    return jsonify({
        'status':     job.get('status'),
        'error':      job.get('error'),
        'pptx_path':  job.get('pptx_path'),
        'pptx_exists': os.path.exists(job['pptx_path']) if job.get('pptx_path') else False,
        'template_exists': os.path.exists(TEMPLATE_PATH),
        'template_path': TEMPLATE_PATH,
        'jobs_dir': JOBS_DIR,
        'sections_count': len(job.get('sections', [])),
        'progress_last': job.get('progress', [{}])[-1] if job.get('progress') else None,
    })

@app.route('/status/<job_id>')
def status(job_id):
    job = load_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    meta = job.get('meta', {})
    return jsonify({
        'status':       job.get('status'),
        'progress':     job.get('progress', []),
        'sections':     job.get('sections', []),
        'meta':         meta,
        'intel':        job.get('intel', {}),
        'error':        job.get('error'),
        'pptx_error':   job.get('pptx_error'),
        'pptx_ready':   bool(job.get('pptx_path') and os.path.exists(job.get('pptx_path',''))),
        # Review panel fields — consumed when status is awaiting_review
        'review': {
            'project_type': meta.get('project_type', ''),
            'client':       meta.get('client', ''),
            'sector':       meta.get('sector', ''),
            'scope_plain':  meta.get('scope_plain', meta.get('scope', '')),
            'brief_summary': meta.get('brief_summary', ''),
        }
    })

@app.route('/rebuild', methods=['POST'])
def rebuild():
    """Rebuild PPTX from edited sections."""
    data = request.get_json()
    job_id = data.get('job_id')
    sections = data.get('sections', [])
    meta = data.get('meta', {})

    if not sections:
        return jsonify({'error': 'No sections provided'}), 400

    try:
        import tempfile
        pptx_dir = tempfile.mkdtemp(prefix='2020_out_')
        pptx_path = os.path.join(pptx_dir, 'proposal.pptx')
        build_pptx_clean(sections, meta, pptx_path)
        if job_id:
            update_job(job_id, pptx_path=pptx_path, status='done', pptx_error=None)
        return jsonify({'status': 'ok', 'job_id': job_id})
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()[-500:]}), 500

@app.route('/download-docx/<job_id>')
def download_docx(job_id):
    job = load_job(job_id)
    if not job:
        return 'Job not found', 404
    sections = job.get('sections', [])
    meta     = job.get('meta', {})
    if not sections:
        return 'No content yet', 400
    try:
        path = build_docx(sections, meta)
        slug = re.sub(r'[^a-zA-Z0-9]+', '_', meta.get('venue', 'Proposal'))
        return send_file(path, as_attachment=True,
                         download_name=f'{slug}_20.20_Proposal.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return f'Error building Word doc: {e}', 500


@app.route('/download/<job_id>')
def download(job_id):
    job = load_job(job_id)
    if not job:
        return 'Job not found — jobs are cleared when the server restarts. Please generate again.', 404
    if job.get('error'):
        return f'Generation failed: {job["error"]}', 500
    if not job.get('pptx_path'):
        return f'PowerPoint not ready yet — status is {job.get("status","unknown")}. Try again in a moment.', 404
    if not os.path.exists(job['pptx_path']):
        return 'PowerPoint file missing — server may have restarted. Please generate again.', 404

    venue = job.get('meta', {}).get('venue', 'Proposal').replace(' ', '_').replace("'", '').replace('&','and').replace('(','').replace(')','')
    filename = f'{venue}_20.20_Proposal.pptx'

    return send_file(
        job['pptx_path'],
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation'
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
